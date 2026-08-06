# ruff: noqa: E501, RUF001
"""Generate the Phase 2 research report and recruiter-facing case study."""

from __future__ import annotations

import json
from pathlib import Path

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
REPORTS = ROOT / "reports" / "generated"
PUBLICATION = ROOT / "reports" / "publication"
FIGURES = ROOT / "reports" / "figures" / "phase2_publication"

NAVY = "17365D"
BLUE = "2E75B6"
PALE = "EAF1F8"
GRAY = "667085"
AMBER = "B66A00"
RED = "B42318"
GREEN = "067647"


def font(run, size=10.5, bold=False, color="202124", italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tc_pr.append(node)
    node.set(qn("w:fill"), fill)


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for cell, width in zip(row.cells, widths, strict=True):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            margins = OxmlElement("w:tcMar")
            for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
                item = OxmlElement(f"w:{side}")
                item.set(qn("w:w"), str(value))
                item.set(qn("w:type"), "dxa")
                margins.append(item)
            tc_pr.append(margins)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, value in enumerate(headers):
        shade(table.rows[0].cells[index], PALE)
        paragraph = table.rows[0].cells[index].paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        font(paragraph.add_run(str(value)), 9, True, NAVY)
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            font(paragraph.add_run(str(value)), 8.7)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def body(doc, text, lead=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.10
    if lead and text.startswith(lead):
        font(paragraph.add_run(lead), bold=True, color=NAVY)
        text = text[len(lead) :]
    font(paragraph.add_run(text))


def bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.167
    font(paragraph.add_run(text), 10.3)


def heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    font(paragraph.add_run(text), {1: 16, 2: 13, 3: 12}[level], True, NAVY)


def figure(doc, path, caption, width=6.25):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    picture = paragraph.add_run().add_picture(str(path), width=Inches(width))
    picture._inline.docPr.set("descr", caption)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    font(cap.add_run(caption), 8.5, italic=True, color=GRAY)


def configure(doc, running_title):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = section.right_margin = Inches(1)
    section.top_margin = section.bottom_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for level, size, before, after in [(1, 16, 16, 8), (2, 13, 12, 6), (3, 12, 8, 4)]:
        style = styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(NAVY if level < 3 else BLUE)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    font(header.add_run(running_title), 8.5, color=GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(footer.add_run("Phase 2 | Development evidence—not a final future test"), 8, color=GRAY)


def cover(doc, kicker, title, subtitle):
    for _ in range(5):
        doc.add_paragraph()
    kicker_p = doc.add_paragraph()
    kicker_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(kicker_p.add_run(kicker.upper()), 10, True, AMBER)
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(8)
    font(title_p.add_run(title), 28, True, NAVY)
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(26)
    font(subtitle_p.add_run(subtitle), 14, color=BLUE)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(
        meta.add_run("SEC EDGAR + FRED | 117 active issuers | Financial cutoff 31 Dec 2025"),
        10,
        color=GRAY,
    )
    doc.add_page_break()


def make_figures(horizon, policy, features):
    FIGURES.mkdir(parents=True, exist_ok=True)
    overall = horizon.loc[horizon["sector"].eq("Overall")].set_index("horizon_quarters")
    fig, axes = plt.subplots(1, 3, figsize=(8.3, 3.6))
    specs = [
        ("PR_AUC", "PR-AUC"),
        ("precision_at_80pct_policy", "Precision"),
        ("alert_rate_at_80pct_policy", "Alert rate"),
    ]
    for axis, (column, title) in zip(axes, specs, strict=True):
        values = overall.loc[[4, 2], column]
        axis.bar(["4 quarters", "2 quarters"], values, color=[f"#{NAVY}", f"#{AMBER}"])
        axis.set_title(title)
        axis.set_ylim(0, max(values.max() * 1.25, 0.35))
        axis.tick_params(axis="x", labelsize=8)
        for index, value in enumerate(values):
            axis.text(index, value + 0.01, f"{value:.1%}", ha="center", fontsize=9)
    fig.suptitle("Paired calibrated horizon comparison", fontweight="bold")
    fig.tight_layout()
    fig.savefig(FIGURES / "horizon_comparison.png", dpi=220)
    plt.close(fig)

    selected = features.loc[features["selected"].eq(True)]
    stability = selected.groupby(["horizon_quarters", "feature"])["fold_id"].nunique().reset_index()
    stability = stability.sort_values(["horizon_quarters", "fold_id"], ascending=[True, False])
    labels = sorted(stability["feature"].unique())
    matrix = (
        stability.pivot(index="feature", columns="horizon_quarters", values="fold_id")
        .fillna(0)
        .reindex(labels)
    )
    fig, axis = plt.subplots(figsize=(8.3, 5.2))
    image = axis.imshow(matrix[[4, 2]], cmap="Blues", vmin=0, vmax=2, aspect="auto")
    axis.set_xticks([0, 1], ["4-quarter", "2-quarter"])
    axis.set_yticks(range(len(matrix)), [x.replace("_", " ") for x in matrix.index], fontsize=7.5)
    axis.set_title("Feature-selection frequency across two temporal folds")
    for row in range(len(matrix)):
        for column in range(2):
            axis.text(
                column, row, int(matrix.iloc[row, column]), ha="center", va="center", fontsize=8
            )
    fig.colorbar(image, ax=axis, label="Selected folds")
    fig.tight_layout()
    fig.savefig(FIGURES / "feature_stability.png", dpi=220)
    plt.close(fig)


def pct(value):
    return f"{float(value):.1%}"


def build_research_report(
    horizon, eligibility, differences, metrics, cases, feature_evidence, policy
):
    doc = Document()
    configure(doc, "Corporate Financial Deterioration | Phase 2 Research Report")
    cover(
        doc,
        "Financial Risk & Data Science Research",
        "Corporate Financial Deterioration Early-Warning Platform",
        "Phase 2: model improvement, interpretability, and horizon sensitivity",
    )
    overall = horizon.loc[horizon["sector"].eq("Overall")].set_index("horizon_quarters")
    sectors = horizon.loc[~horizon["sector"].eq("Overall")].set_index(
        ["horizon_quarters", "sector"]
    )

    heading(doc, "Executive summary")
    body(
        doc,
        "This study develops a point-in-time analyst-screening system for detecting weakening corporate debt-service capacity. The system is not a bankruptcy model and does not replace credit analysis. It prioritizes recall so that potentially deteriorating firms reach manual review.",
    )
    body(
        doc,
        "Main conclusion. The two-quarter sensitivity model met the 80% recall requirement in both sectors and reduced the paired alert rate by 8.3 percentage points. However, PR-AUC declined by 0.090 and precision declined by 6.6 percentage points. The shorter horizon therefore reduces queue size, but the remaining alerts are less concentrated with true events. It is useful as a secondary near-term screen, not a superior replacement for the four-quarter primary outcome.",
        "Main conclusion. ",
    )
    add_table(
        doc,
        ["Paired calibrated result", "4-quarter primary", "2-quarter sensitivity"],
        [
            [
                "Eligible validation rows",
                f"{int(overall.loc[4, 'evaluated_company_quarters']):,}",
                f"{int(overall.loc[2, 'evaluated_company_quarters']):,}",
            ],
            [
                "Event prevalence",
                pct(overall.loc[4, "event_prevalence"]),
                pct(overall.loc[2, "event_prevalence"]),
            ],
            ["PR-AUC", f"{overall.loc[4, 'PR_AUC']:.3f}", f"{overall.loc[2, 'PR_AUC']:.3f}"],
            [
                "Recall",
                pct(overall.loc[4, "recall_at_80pct_policy"]),
                pct(overall.loc[2, "recall_at_80pct_policy"]),
            ],
            [
                "Precision",
                pct(overall.loc[4, "precision_at_80pct_policy"]),
                pct(overall.loc[2, "precision_at_80pct_policy"]),
            ],
            [
                "Alert rate",
                pct(overall.loc[4, "alert_rate_at_80pct_policy"]),
                pct(overall.loc[2, "alert_rate_at_80pct_policy"]),
            ],
            [
                "Brier score",
                f"{overall.loc[4, 'Brier_score']:.3f}",
                f"{overall.loc[2, 'Brier_score']:.3f}",
            ],
            [
                "Median warning lead",
                f"{overall.loc[4, 'median_warning_lead_quarters']:.1f} quarters",
                f"{overall.loc[2, 'median_warning_lead_quarters']:.1f} quarters",
            ],
        ],
        [3600, 2880, 2880],
    )
    figure(
        doc,
        FIGURES / "horizon_comparison.png",
        "Figure 1. The shorter horizon lowers workload but also weakens ranking and alert precision.",
    )

    doc.add_page_break()
    heading(doc, "1. Business and financial question")
    body(
        doc,
        "Credit and equity analysts often screen many issuers before performing detailed filing review, covenant analysis, valuation work, or management assessment. The project asks whether public financial statements and macroeconomic information can rank active companies by the risk of a material decline in interest coverage.",
    )
    bullet(
        doc,
        "Primary financial outcome: within four future fiscal quarters, interest coverage falls below 1.5x and declines at least 40% from the current level.",
    )
    bullet(
        doc,
        "Secondary sensitivity outcome: apply the same economic rule over two future fiscal quarters.",
    )
    bullet(
        doc,
        "Decision objective: attain at least 80% recall separately in Consumer Discretionary and Utilities, then minimize the number of alerts.",
    )
    body(
        doc,
        "Interest coverage connects operating earnings to contractual interest expense. A decline below 1.5x indicates a narrow earnings cushion, but the rule is a research definition—not a covenant, default forecast, or investment recommendation.",
    )

    heading(doc, "2. Data, population, and sampling")
    body(
        doc,
        "The analysis uses SEC EDGAR company facts and filing metadata for financial statements, plus FRED/ALFRED macroeconomic series. Every field carries an availability date so that a decision row contains only information public at that time.",
    )
    add_table(
        doc,
        ["Design element", "Phase 2 treatment"],
        [
            [
                "Population",
                "Currently listed US operating companies; delisted companies excluded by scope",
            ],
            ["Selection date", "2 August 2026"],
            ["Financial cutoff", "31 December 2025"],
            ["Sample", "75 Consumer Discretionary and 42 Utilities issuers"],
            ["Randomization", "Seeded stratified random sampling without replacement; seed 42"],
            [
                "Eligibility",
                "Reliable interest coverage and assets; occasional optional-field gaps permitted",
            ],
            ["Synthetic data", "None"],
        ],
        [2500, 6860],
    )
    body(
        doc,
        "Grocery-dominant discount retailers are excluded because their sales mix is staples-oriented. Sector assignment uses SEC SIC mappings and reviewed business descriptions; store format alone does not determine sector membership.",
    )

    heading(doc, "3. Point-in-time analytical pipeline")
    for item in [
        "Retrieve and checksum public source responses without committing credentials or raw data.",
        "Normalize XBRL concepts, construct financial ratios, and preserve filing and availability lineage.",
        "Join only macro vintages available at each decision date.",
        "Build future labels only where all required future quarters are consecutive and observed.",
    ]:
        bullet(doc, item)

    heading(doc, "4. Modeling design and leakage controls")
    body(
        doc,
        "The primary classifier is a regularized partially pooled logistic model. It shares a common relationship across sectors while allowing a small number of Utility deviations. A pooled logistic model is the benchmark and constrained gradient boosting is the nonlinear challenger.",
    )
    heading(doc, "Expanding-window validation", 2)
    body(
        doc,
        "Each validation period occurs after its training history. Labels are admitted to training only when the full outcome window was already observable before the validation origin. This embargo prevents future outcomes from leaking backward into model fitting.",
    )
    heading(doc, "Fold-local feature selection", 2)
    body(
        doc,
        "Within each outer training fold, candidate variables pass missingness and variance checks, correlation pruning, and repeated temporal permutation-importance testing. The outer validation period is never used to choose features. Missing indicators and imputation are also fitted only within training folds.",
    )
    figure(
        doc,
        FIGURES / "feature_stability.png",
        "Figure 2. Counts show how often each candidate survived fold-local selection; zero means it was evaluated but not retained.",
    )
    body(
        doc,
        "Refinancing gap/assets was selected in one of two folds, but was not stable and was not retained by the latest-fold recommendation. Filing delay is excluded from the classifier and appears only as a data-quality and case-review field.",
    )

    heading(doc, "5. Four-quarter primary model results")
    primary_metrics = metrics.loc[metrics["model"].str.startswith("partially_pooled")]
    boost_metrics = metrics.loc[metrics["model"].str.startswith("pooled_gradient")]
    rows = []
    for label, table in [
        ("Partially pooled logistic", primary_metrics),
        ("Constrained boosting", boost_metrics),
    ]:
        overall_row = table.loc[table["slice"].eq("Overall")].iloc[0]
        rows.append(
            [
                label,
                f"{overall_row['PR_AUC']:.3f}",
                f"{overall_row['precision']:.3f}",
                f"{overall_row['recall']:.3f}",
                f"{overall_row['Brier_score']:.3f}",
            ]
        )
    add_table(
        doc,
        ["Model", "PR-AUC", "Precision*", "Recall*", "Brier"],
        rows,
        [3500, 1465, 1465, 1465, 1465],
    )
    body(
        doc,
        "*Precision and recall in this model-comparison table use the common 20% alert-rate reference point. The separate operational threshold policy targets 80% sector recall.",
    )
    body(
        doc,
        "The boosting challenger ranked events better, but the partially pooled logistic model remains the registered primary because its structure and financial relationships are easier to defend. This is an explicit interpretability-performance tradeoff rather than an attempt to hide the stronger challenger.",
    )

    heading(doc, "6. Two-quarter sensitivity experiment")
    body(
        doc,
        "The comparison was frozen before examining results: four quarters remained primary; two quarters used the identical 1.5x and 40% rules; both horizons used the same issuers, same paired validation company-quarters, same fold calendar, and independent fold-local feature selection. Calibration used only earlier out-of-fold predictions.",
    )
    full_rows = []
    for horizon_value in [4, 2]:
        for sector in ["Consumer Discretionary", "Utilities"]:
            row = sectors.loc[(horizon_value, sector)]
            full_rows.append(
                [
                    f"{horizon_value}Q",
                    sector.replace("Consumer Discretionary", "Consumer"),
                    f"{int(row['evaluated_company_quarters']):,}",
                    pct(row["event_prevalence"]),
                    f"{row['PR_AUC']:.3f}",
                    pct(row["recall_at_80pct_policy"]),
                    pct(row["precision_at_80pct_policy"]),
                    pct(row["alert_rate_at_80pct_policy"]),
                ]
            )
    add_table(
        doc,
        ["Horizon", "Sector", "Rows", "Prev.", "PR-AUC", "Recall", "Precision", "Alerts"],
        full_rows,
        [850, 1800, 900, 1050, 1100, 1100, 1260, 1300],
    )
    body(
        doc,
        "Both horizons exceeded 80% recall in each sector. Utilities required much larger review queues and produced lower precision at both horizons, reflecting a smaller event set and weaker ranking. Equal recall is not forced; each sector simply has to clear the minimum standard.",
    )

    heading(doc, "Statistical uncertainty", 2)
    diff_rows = []
    names = {
        "PR_AUC": "PR-AUC",
        "precision": "Precision",
        "alert_rate": "Alert rate",
        "Brier_score": "Brier score",
        "median_warning_lead_quarters": "Median lead (quarters)",
    }
    for key, name in names.items():
        row = differences.loc[differences["metric"].eq(key)].iloc[0]
        diff_rows.append(
            [
                name,
                f"{row['difference_2q_minus_4q']:.3f}",
                f"[{row['clustered_95pct_lower']:.3f}, {row['clustered_95pct_upper']:.3f}]",
            ]
        )
    add_table(
        doc,
        ["Two-quarter minus four-quarter", "Estimate", "Company-clustered 95% interval"],
        diff_rows,
        [4200, 1900, 3260],
    )
    body(
        doc,
        "The alert-rate reduction, precision decline, and PR-AUC decline all have company-clustered intervals that exclude zero. Companies—not individual quarters—were resampled because observations from the same issuer are correlated.",
    )

    heading(doc, "Calibration and lead time", 2)
    calibration = policy.set_index("horizon_quarters")
    body(
        doc,
        f"The four-quarter paired model used pooled {calibration.loc[4, 'calibration_method']} recalibration and achieved a Brier score of {overall.loc[4, 'Brier_score']:.3f} with expected calibration error {overall.loc[4, 'expected_calibration_error']:.3f}. The two-quarter model retained {calibration.loc[2, 'calibration_method']} calibration; its Brier score was {overall.loc[2, 'Brier_score']:.3f} and calibration error {overall.loc[2, 'expected_calibration_error']:.3f}.",
    )
    body(
        doc,
        "The lower two-quarter Brier score partly reflects lower event prevalence, so it should not be read as proof of better ranking. Median warning lead fell from 2.0 to 1.0 quarters in the correctly alerted paired events. That is structurally expected: a shorter outcome window cannot capture later-onset deterioration.",
    )

    doc.add_page_break()
    heading(doc, "7. Company-level financial case analysis")
    for _, case in cases.iterrows():
        title = case["case_type"].replace("_", " ").title()
        heading(doc, f"{title}: {case['company_name']} ({case['ticker']})", 2)
        body(
            doc,
            f"At {pd.Timestamp(case['decision_at']).date()}, the four-quarter model scored {case['probability']:.1%} against its {case['sector_threshold']:.1%} sector threshold. Interest coverage was {case['interest_coverage_ttm']:.2f}x, free-cash-flow margin {case['free_cash_flow_margin_ttm']:.1%}, operating margin {case['operating_margin_ttm']:.1%}, and debt/assets {case['total_debt_to_assets']:.1%}.",
        )
        reasons = json.loads(case["reason_codes_json"])
        if reasons:
            body(
                doc,
                "Associated model-facing conditions included "
                + "; ".join(reason["label"].lower() for reason in reasons)
                + ". These are predictive associations, not causal findings.",
            )
        else:
            body(
                doc,
                "No monitored ratio was extreme enough to generate a standard reason code. This absence is itself informative: the later decline was not apparent in the model's strongest contemporaneous indicators.",
            )
        body(doc, case["analyst_usefulness"], "")
        body(doc, case["result_explanation"])

    heading(doc, "8. Financial interpretation")
    bullet(
        doc, "The model is strongest as a prioritization system, not an autonomous credit decision."
    )
    bullet(
        doc,
        "Weak peer-relative cash generation and debt-service capacity are understandable review triggers.",
    )
    bullet(
        doc,
        "A false alert can still identify a financially weak company even when the exact future label is not met.",
    )
    bullet(
        doc,
        "A missed event can arise when current financial ratios appear healthy but deteriorate later; no historical classifier can observe an unreported future shock.",
    )

    doc.add_page_break()
    heading(doc, "9. Limitations and model governance")
    for text in [
        "The active-company-only sample creates survivorship bias and cannot represent delisted or failed firms.",
        "The 2023-and-later Phase 1 period has already been examined, so Phase 2 results are development evidence rather than a final untouched test.",
        "Only two sectors are included, and Utilities have fewer issuers and deterioration events.",
        "Overlapping quarterly outcomes are correlated; company-clustered uncertainty reduces but does not eliminate this concern.",
        "Thresholds were estimated from development predictions and may require recalibration when future outcomes mature.",
        "The label measures a specific interest-coverage deterioration pattern, not default, downgrade, fraud, or investment loss.",
    ]:
        bullet(doc, text)
    heading(doc, "Readiness conclusion", 2)
    body(
        doc,
        "The implementation objectives are complete, but Phase 2 is not ready for a definitive future-performance claim. A new test beginning after the frozen 2025 financial cutoff should remain unopened until four-quarter outcomes mature. Until then, all reported values should be described as temporal out-of-fold development evidence.",
    )

    heading(doc, "10. Reproducibility and data lineage")
    body(
        doc,
        "The repository provides a locked Python environment, versioned YAML configuration, deterministic seeds, ordered CLI stages, automated tests, source manifests, model/data cards, and generated evidence tables. Private SEC identification and the FRED key belong only in the ignored .env file. The committed .env.example contains placeholders.",
    )
    add_table(
        doc,
        ["Artifact", "Purpose"],
        [
            [
                "configs/phase2.yml",
                "Frozen sampling, label, validation, model, and threshold policy",
            ],
            ["docs/phase2_methodology.md", "Technical definitions in plain language"],
            [
                "docs/phase2_reproducibility.md",
                "Environment, commands, source attribution, and expected outputs",
            ],
            ["reports/generated/phase2_horizon_*.csv", "Auditable empirical sensitivity results"],
            [
                "reports/generated/phase2_company_case_studies.csv",
                "Selected TP, FP, and missed-event evidence",
            ],
            ["tests/", "Unit, leakage, selection, sampling, and policy checks"],
        ],
        [3600, 5760],
    )

    heading(doc, "11. Conclusion")
    body(
        doc,
        "The project demonstrates a complete financial data-science workflow: point-in-time data engineering, economically interpretable labels and features, imbalance-aware evaluation, temporal cross-validation, fold-local preprocessing and feature selection, calibration, threshold optimization, uncertainty analysis, and case-level interpretation.",
    )
    body(
        doc,
        "The empirical horizon test does not support replacing the four-quarter primary outcome. Two quarters reduces workload and produces probabilities with a lower Brier score, but it identifies a narrower and more immediate event, ranks those events less effectively, lowers precision, and cuts warning lead time. The defensible deployment concept is therefore a two-layer analyst screen: four quarters for medium-term deterioration and two quarters as an optional near-term sensitivity flag.",
    )
    return doc


def build_case_study(horizon, cases):
    doc = Document()
    configure(doc, "Portfolio Case Study | Financial Deterioration Screening")
    cover(
        doc,
        "Data Science Portfolio Case Study",
        "Financial Deterioration Screening",
        "From point-in-time SEC data to an analyst review queue",
    )
    overall = horizon.loc[horizon["sector"].eq("Overall")].set_index("horizon_quarters")
    heading(doc, "The business problem")
    body(
        doc,
        "An analyst cannot deeply review every public company every quarter. This project builds a screening layer that ranks companies for follow-up when operating earnings may become insufficient relative to interest expense. False alerts are reviewed manually; missed deterioration is considered more costly.",
    )
    heading(doc, "What I built")
    for text in [
        "A filing-aware SEC EDGAR and FRED/ALFRED pipeline with source and availability lineage.",
        "A reproducible 117-company sample: 75 Consumer Discretionary and 42 Utilities issuers.",
        "Financial-ratio, trend, volatility, and peer-relative features with fold-local missing-data treatment.",
        "Expanding-window model validation with outcome embargoes and company-clustered uncertainty.",
        "An interpretable partially pooled logistic primary, pooled benchmark, and constrained boosting challenger.",
        "Sector-specific thresholds that attain at least 80% recall while minimizing review workload.",
    ]:
        bullet(doc, text)
    doc.add_page_break()
    heading(doc, "Point-in-time pipeline and leakage controls")
    body(
        doc,
        "The most important engineering requirement was to reproduce what an analyst could actually have known at each historical decision date. Fiscal period-end dates were not treated as publication dates. SEC filing availability, amendments, macro vintages, and feature availability were retained explicitly.",
    )
    add_table(
        doc,
        ["Leakage risk", "Control"],
        [
            ["Using later filings", "Filter every fact by its public availability date"],
            [
                "Using immature outcomes",
                "Require the complete consecutive future window before labeling",
            ],
            [
                "Global imputation",
                "Fit imputation and missing indicators inside each training fold",
            ],
            [
                "Feature-selection leakage",
                "Select variables only within outer-fold training history",
            ],
            ["Random train/test mixing", "Use expanding calendar windows and label embargoes"],
            [
                "Threshold overstatement",
                "Label all metrics as development OOF until a new future test matures",
            ],
        ],
        [2800, 6560],
    )
    doc.add_page_break()
    heading(doc, "Modeling and feature decisions")
    body(
        doc,
        "The primary logistic model was chosen for interpretability, not because it had the highest score. Gradient boosting reached 0.412 development PR-AUC versus 0.379 for the partially pooled logistic model. The challenger is stronger statistically, while the primary is easier to explain to an analyst.",
    )
    body(
        doc,
        "Stable selected variables included current interest coverage, its year-over-year change, volatility and sector percentile, free-cash-flow-margin sector percentile, operating margin, and revenue growth. Refinancing gap was selected in one early fold but was unstable and absent from the latest-fold recommendation; filing delay remained a data-quality field rather than direct distress evidence.",
    )
    figure(
        doc,
        FIGURES / "feature_stability.png",
        "Fold-local selection emphasizes understandable operating, cash-flow, and debt-service measures.",
    )
    doc.add_page_break()
    heading(doc, "Main operating result")
    body(
        doc,
        f"At the paired calibrated four-quarter policy, overall recall was {overall.loc[4, 'recall_at_80pct_policy']:.1%}, precision was {overall.loc[4, 'precision_at_80pct_policy']:.1%}, and {overall.loc[4, 'alert_rate_at_80pct_policy']:.1%} of company-quarters entered review. High recall is achievable, but the workload is substantial.",
    )
    figure(
        doc,
        FIGURES / "horizon_comparison.png",
        "The two-quarter sensitivity reduces workload but weakens ranking and precision.",
    )
    body(
        doc,
        f"The two-quarter sensitivity retained {overall.loc[2, 'recall_at_80pct_policy']:.1%} recall and reduced alerts to {overall.loc[2, 'alert_rate_at_80pct_policy']:.1%}. Its PR-AUC fell from {overall.loc[4, 'PR_AUC']:.3f} to {overall.loc[2, 'PR_AUC']:.3f}, however, and precision fell from {overall.loc[4, 'precision_at_80pct_policy']:.1%} to {overall.loc[2, 'precision_at_80pct_policy']:.1%}.",
    )
    doc.add_page_break()
    heading(doc, "Three real model decisions")
    for _, case in cases.iterrows():
        heading(doc, f"{case['case_type'].replace('_', ' ').title()}: {case['ticker']}", 2)
        body(
            doc,
            f"The model assigned {case['probability']:.1%} against a {case['sector_threshold']:.1%} threshold. Current interest coverage was {case['interest_coverage_ttm']:.2f}x and future minimum coverage was {case['future_minimum_interest_coverage']:.2f}x.",
        )
        body(doc, case["analyst_usefulness"])
        body(doc, case["result_explanation"])
    doc.add_page_break()
    heading(doc, "Limitations and what I would do next")
    for text in [
        "Keep the post-2025 future period unopened until four-quarter outcomes mature.",
        "Measure real analyst capacity before operationalizing a 60%–70% alert rate.",
        "Recalibrate probabilities when new labels mature; investigate drift rather than retraining automatically.",
        "Consider adding Industrials only after validating sector mappings and KPI comparability.",
        "Preserve the four-quarter outcome as primary and use two quarters only as a secondary near-term screen.",
    ]:
        bullet(doc, text)
    heading(doc, "Skills demonstrated")
    body(
        doc,
        "Financial statement analysis, SEC/FRED data engineering, point-in-time joins, temporal cross-validation, missing-data methods, imbalanced classification, feature selection, calibration, interpretability, clustered uncertainty, model governance, automated testing, and communication of statistical results in a financial decision context.",
    )
    doc.add_page_break()
    heading(doc, "Resume-ready summary")
    bullet(
        doc,
        "Developed a point-in-time financial-deterioration screening pipeline using SEC filings and macro data for 117 active Consumer Discretionary and Utility issuers.",
    )
    bullet(
        doc,
        "Implemented expanding-window validation, fold-local feature selection, calibrated classification, and sector thresholds exceeding 80% recall in both sectors.",
    )
    bullet(
        doc,
        "Compared an interpretable partially pooled logistic model with pooled logistic and constrained boosting; the challenger achieved 0.412 development PR-AUC.",
    )
    bullet(
        doc,
        "Empirically tested two- versus four-quarter warning horizons using paired folds and issuer-clustered bootstrap uncertainty, documenting a workload-versus-precision tradeoff.",
    )
    body(
        doc,
        "All results are explicitly labeled as development evidence. The project does not claim bankruptcy prediction, causal effects, or final prospective performance.",
    )
    return doc


def main():
    PUBLICATION.mkdir(parents=True, exist_ok=True)
    horizon = pd.read_csv(REPORTS / "phase2_horizon_comparison.csv")
    eligibility = pd.read_csv(REPORTS / "phase2_horizon_eligibility_and_prevalence.csv")
    differences = pd.read_csv(REPORTS / "phase2_horizon_clustered_differences.csv")
    metrics = pd.read_csv(REPORTS / "phase2_metrics.csv")
    cases = pd.read_csv(REPORTS / "phase2_company_case_studies.csv")
    feature_evidence = pd.read_csv(REPORTS / "phase2_horizon_feature_selection_evidence.csv")
    policy = pd.read_csv(REPORTS / "phase2_horizon_policy_comparison.csv")
    make_figures(horizon, policy, feature_evidence)
    research = build_research_report(
        horizon, eligibility, differences, metrics, cases, feature_evidence, policy
    )
    case_study = build_case_study(horizon, cases)
    research.save(PUBLICATION / "Corporate_Financial_Deterioration_Phase2_Research_Report.docx")
    case_study.save(
        PUBLICATION / "Corporate_Financial_Deterioration_Phase2_Portfolio_Case_Study.docx"
    )


if __name__ == "__main__":
    main()
