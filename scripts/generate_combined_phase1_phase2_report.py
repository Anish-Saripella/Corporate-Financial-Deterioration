# ruff: noqa: E501
"""Generate the integrated Phase 1 + Phase 2 research report."""

from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from generate_phase2_publication import (
    AMBER,
    GRAY,
    NAVY,
    add_table,
    body,
    bullet,
    configure,
    cover,
    figure,
    font,
    heading,
)

ROOT = Path(__file__).resolve().parents[1]
GENERATED = ROOT / "reports" / "generated"
PUBLICATION = ROOT / "reports" / "publication"
FIGURES = ROOT / "reports" / "figures" / "combined_publication"
OUTPUT = (
    PUBLICATION / "Corporate_Financial_Deterioration_Combined_Phase1_Phase2_Research_Report.docx"
)


def add_centered_page_number(paragraph) -> None:
    """Add a Word PAGE field to the centered footer paragraph."""
    paragraph.text = ""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    font(run, 8, color=GRAY)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instruction, separate, value, end):
        run._r.append(element)


def make_figures(horizon: pd.DataFrame) -> None:
    import matplotlib.pyplot as plt

    FIGURES.mkdir(parents=True, exist_ok=True)

    fig, ax = plt.subplots(figsize=(8.2, 3.8))
    labels = ["Phase 1\nlocked holdout", "Phase 2\npartial pooling", "Phase 2\nboosting challenger"]
    values = [0.397, 0.379378, 0.412197]
    bars = ax.bar(labels, values, color=[f"#{GRAY}", f"#{NAVY}", f"#{AMBER}"])
    ax.set_ylim(0, 0.50)
    ax.set_ylabel("PR-AUC")
    ax.set_title("Classification ranking evidence by research stage")
    ax.grid(axis="y", alpha=0.2)
    for bar, value in zip(bars, values, strict=True):
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            value + 0.012,
            f"{value:.3f}",
            ha="center",
            fontweight="bold",
        )
    ax.text(
        0.5,
        -0.24,
        "Evidence sets differ: Phase 1 is final holdout; Phase 2 is temporal development OOF.",
        transform=ax.transAxes,
        ha="center",
        fontsize=9,
        color=f"#{GRAY}",
    )
    fig.tight_layout()
    fig.savefig(FIGURES / "performance_evidence.png", dpi=220, bbox_inches="tight")
    plt.close(fig)

    overall = horizon.loc[horizon["sector"].eq("Overall")].set_index("horizon_quarters")
    fig, axes = plt.subplots(1, 3, figsize=(8.2, 3.7))
    for ax, column, title in [
        (axes[0], "PR_AUC", "PR-AUC"),
        (axes[1], "precision_at_80pct_policy", "Precision"),
        (axes[2], "median_warning_lead_quarters", "Median lead (quarters)"),
    ]:
        vals = [overall.loc[4, column], overall.loc[2, column]]
        bars = ax.bar(["4Q", "2Q"], vals, color=[f"#{NAVY}", f"#{AMBER}"])
        ax.set_title(title)
        ax.grid(axis="y", alpha=0.2)
        ax.set_ylim(0, max(vals) * 1.3)
        for bar, value in zip(bars, vals, strict=True):
            label = (
                f"{value:.3f}"
                if title == "PR-AUC"
                else (f"{value:.1%}" if title == "Precision" else f"{value:.1f}")
            )
            ax.text(
                bar.get_x() + bar.get_width() / 2,
                value + max(vals) * 0.04,
                label,
                ha="center",
                fontsize=9,
            )
    fig.suptitle(
        "Paired horizon sensitivity: shorter is not automatically better", fontweight="bold"
    )
    fig.tight_layout()
    fig.savefig(FIGURES / "horizon_tradeoff.png", dpi=220)
    plt.close(fig)


def add_evidence_box(doc: Document) -> None:
    add_table(
        doc,
        ["Evidence level", "Status", "Correct interpretation"],
        [
            [
                "1. Phase 1 final evaluation",
                "2023+ locked holdout opened once",
                "Final benchmark for the original 60-company experiment",
            ],
            [
                "2. Phase 2 initial model development",
                "Expanding-window OOF",
                "Model-selection and policy evidence; not a final future test",
            ],
            [
                "3. Phase 2 model optimization",
                "Sealed late-2024 out-of-time test",
                "Final test evidence for the frozen 117-company optimization experiment",
            ],
            [
                "4. Two-quarter sensitivity",
                "Paired OOF experiment",
                "Tests a different warning horizon; does not replace the primary label",
            ],
            [
                "5. Post-2025 prospective test",
                "Unopened and outcome-immature",
                "Provides an additional future-period replication after outcomes mature",
            ],
        ],
        [2500, 2450, 4410],
    )


def build_report(*, regenerate_figures: bool = True) -> Document:
    horizon = pd.read_csv(GENERATED / "phase2_horizon_comparison.csv")
    cases = pd.read_csv(GENERATED / "phase2_company_case_studies.csv")
    recall_policy = pd.read_csv(GENERATED / "phase2_recall_first_threshold_table.csv")
    optimized = json.loads(
        (ROOT / "reports" / "source_data" / "phase3_final_metrics.json").read_text(encoding="utf-8")
    )
    if regenerate_figures:
        make_figures(horizon)

    doc = Document()
    configure(
        doc, "Corporate Financial Deterioration | Integrated Financial Risk Research Analysis"
    )
    add_centered_page_number(doc.sections[0].footer.paragraphs[0])
    cover(
        doc,
        "Financial Risk & Data Science Research",
        "Corporate Financial Deterioration Early-Warning Platform",
        "Integrated Financial Risk Research Analysis",
        meta_text=False,
    )

    heading(doc, "Executive summary")
    body(
        doc,
        "This project asks whether public SEC filings and historically available macroeconomic data can identify active U.S. companies whose debt-service capacity may deteriorate over the next four fiscal quarters. The product is a ranked review queue for analysts. It is not a bankruptcy model, credit rating, valuation model, or automated investment recommendation.",
    )
    body(
        doc,
        "What Phase 1 established. A point-in-time pipeline, financially interpretable deterioration label, time-series forecasts, and imbalanced classifier were tested on 60 active issuers. The frozen gradient-boosted model achieved 0.397 PR-AUC, 56.3% recall, 33.3% precision, 1.97x top-decile lift, and 0.159 Brier score on 457 locked 2023-and-later observations.",
        "What Phase 1 established. ",
    )
    body(
        doc,
        "What Phase 2 improved. The sample increased to 117 active issuers; eligibility permitted defensible optional-field gaps; feature selection moved inside temporal folds; the registered model became an interpretable partially pooled logistic specification; calibration, sector thresholds, clustered uncertainty, monitoring, and company-level explanations were added. The primary reached 0.379 development PR-AUC and the constrained boosting challenger reached 0.412.",
        "What Phase 2 improved. ",
    )
    body(
        doc,
        "What subsequent Phase 2 optimization achieved. The same population, data, and four-quarter outcome were retained while additional model families and ensemble techniques were compared through rolling temporal validation. A frozen blend of pooled and sector-specific XGBoost achieved 0.841 ROC-AUC, 0.494 PR-AUC, and 85.7% recall on a sealed late-2024 out-of-time test whose outcomes matured during 2025.",
        "What subsequent Phase 2 optimization achieved. ",
    )
    body(
        doc,
        "Decision conclusion. The project has achieved a credible research prototype for recall-first analyst triage and a successful out-of-time model evaluation. Four quarters should remain the primary horizon. Two quarters is useful only as a near-term secondary flag because it reduced workload but also reduced PR-AUC, precision, and warning lead. A later matured period is still needed to determine whether the improvement repeats over time.",
        "Decision conclusion. ",
    )
    add_evidence_box(doc)

    doc.add_page_break()
    heading(doc, "1. Financial and business problem")
    body(
        doc,
        "A credit, equity, or risk analyst cannot manually conduct a full filing, covenant, cash-flow, and management-quality review for every issuer each quarter. A screening model can reduce this search problem by ranking firms whose operating earnings may become insufficient relative to their interest burden. Because the model sits before manual analysis, the preferred error trade-off is recall-first: missing a true deterioration is considered more costly than reviewing a false alert.",
    )
    body(
        doc,
        "The outcome is deliberately narrower than default. A company is labeled as deteriorating when future interest coverage both falls below 1.5x and declines at least 40% from the current level. The level condition identifies a thin earnings cushion; the relative decline condition avoids labeling a small change from an already strong position. It does not encode covenant terms, maturity schedules, liquidity facilities, asset values, or recovery rates.",
    )
    add_table(
        doc,
        ["Stakeholder question", "Analytical output", "What still requires human judgment"],
        [
            [
                "Which issuers deserve review first?",
                "Sector-aware probability and alert queue",
                "Material one-time items, strategy, regulation, covenant detail",
            ],
            [
                "Why was a company flagged?",
                "Observed KPI, trend, peer-position, and model reason codes",
                "Causal explanation and management intent",
            ],
            [
                "How uncertain is the outlook?",
                "Forecast ranges, calibration measures, and model stability",
                "Scenario severity and credit/investment action",
            ],
        ],
        [2700, 3000, 3660],
    )

    heading(doc, "2. SEC/FRED data architecture")
    body(
        doc,
        "SEC EDGAR Company Facts supplies standardized XBRL statement facts, filing dates, fiscal periods, and amendments. FRED/ALFRED supplies interest-rate and business-cycle series with historical vintages. The architecture stores raw-response checksums and metadata, maps issuer-specific XBRL concepts into common financial fields, aligns company fiscal quarters, and then builds a company-quarter panel.",
    )
    bullet(
        doc,
        "Raw acquisition: authenticated SEC and FRED requests, retrieval timestamps, source URLs, checksums, and failure logs.",
    )
    bullet(
        doc,
        "Normalization: equivalent accounting tags are mapped to consistent economic concepts and validated for duration, units, sign, fiscal period, and duplicate filings.",
    )
    bullet(
        doc,
        "Point-in-time assembly: a fact or macro vintage enters a decision row only after it was publicly available.",
    )
    bullet(
        doc,
        "Analytical layers: KPIs, forecasts, features, labels, temporal folds, calibrated probabilities, sector thresholds, explanations, and monitoring tables.",
    )
    body(
        doc,
        "No synthetic data are used. Missing observations remain real missing observations until a fold-local model transformation handles them; the pipeline does not invent company histories.",
    )

    doc.add_page_break()
    heading(doc, "3. Population, sector classification, and sampling")
    body(
        doc,
        "The population is currently active U.S. public operating companies at the selection date. Delisted companies remain excluded because the defined research question concerns companies active today. This choice is consistent across phases, but it creates survivorship bias: findings cannot be generalized to the full historical population of failed or delisted issuers.",
    )
    heading(doc, "How a company is classified into a sector", 2)
    body(
        doc,
        "Sector classification begins with the issuer's SEC Standard Industrial Classification code and a documented SIC-to-sector mapping. The mapping is then reviewed against the company's business description and dominant revenue activity. Classification follows the economic source of sales and cash flows—not a superficial label such as store format.",
    )
    add_table(
        doc,
        ["Design", "Phase 1", "Phase 2"],
        [
            [
                "Selection / financial cutoffs",
                "Original frozen Phase 1 snapshot",
                "2 Aug 2026 / 31 Dec 2025",
            ],
            ["Sample", "30 Consumer + 30 Utilities", "75 Consumer + 42 Utilities"],
            [
                "Randomization",
                "Seeded sector-stratified sample and prespecified replacement",
                "Random without replacement, seed 42, within eligible sector pools",
            ],
            [
                "Replacement",
                "14 firms replaced after strict certification",
                "Ranked reserves used only under the same eligibility-based, outcome-blind rules",
            ],
            ["Issuer status", "Active companies only", "Active companies only; delisted excluded"],
            ["Synthetic observations", "None", "None"],
        ],
        [1900, 3420, 4040],
    )
    body(
        doc,
        "Phase 2 ranks all mapped active candidates to create an auditable eligible pool, then applies seeded stratified random sampling without replacement. Randomization reduces researcher discretion; determinism means seed 42 reproduces the exact sample. The 42-company Utilities cap reflects the number that met the revised eligibility requirements. Increasing it to all mapped names would have admitted companies lacking sufficiently reliable core coverage for the defined analysis.",
    )
    heading(doc, "Eligibility change", 2)
    body(
        doc,
        "Phase 1 required complete histories across several fields and lineage checks. Phase 2 keeps strong requirements for interest coverage, assets, quarterly continuity, filing lineage, and point-in-time availability, but allows occasional gaps in optional features such as free-cash-flow margin. This is appropriate because forecasting and classification pipelines can use incomplete histories, provided missingness is handled inside each training fold. It is not appropriate to relax the core outcome or timing fields, because the label and leakage controls would become unreliable.",
    )

    heading(doc, "4. Financial KPI definitions")
    add_table(
        doc,
        ["KPI", "Definition", "Financial interpretation"],
        [
            [
                "Interest coverage",
                "Trailing-12-month operating income / trailing-12-month interest expense",
                "Operating earnings available per dollar of interest burden; lower values mean less debt-service cushion",
            ],
            [
                "Free-cash-flow margin",
                "(TTM operating cash flow - TTM capital expenditure) / TTM revenue",
                "Cash retained after operating and investment needs relative to sales",
            ],
            [
                "Debt-to-assets",
                "Total debt / total assets",
                "Share of the asset base financed by interest-bearing debt; a balance-sheet leverage measure",
            ],
        ],
        [2000, 3500, 3860],
    )
    body(
        doc,
        "Ratios are complementary. Coverage focuses on earnings service capacity, free cash flow on internal funding resilience, and debt-to-assets on structural leverage. Sector context matters: negative free cash flow can be normal for capital-intensive Utilities, while a sharp deterioration from a company's own history may still be important.",
    )

    doc.add_page_break()
    heading(doc, "5. Point-in-time controls and leakage prevention")
    body(
        doc,
        "Financial prediction is especially vulnerable to look-ahead bias because a quarter-end date is not the same as the date investors could see the filing. The pipeline therefore associates each value with its public availability date and admits it only after that date. Amendments are handled through filing lineage rather than silently replacing the historical value with a later revision.",
    )
    add_table(
        doc,
        ["Leakage risk", "Control"],
        [
            ["Using a filing before publication", "Filter SEC facts by filing availability date"],
            [
                "Using revised macro history",
                "Join the FRED/ALFRED vintage available at the decision date",
            ],
            [
                "Training on outcomes not yet observable",
                "Embargo rows until the complete future label window has matured",
            ],
            [
                "Global imputation or scaling",
                "Fit missing-value treatment and transformations within each training fold",
            ],
            [
                "Selecting features on validation data",
                "Run feature selection only on the outer-fold training history",
            ],
            ["Randomly mixing past and future", "Use expanding-window validation ordered by time"],
        ],
        [3350, 6010],
    )
    body(
        doc,
        "Expanding-window validation trains on the available past and validates the model on a later period before expanding the training window. As a result, every validation prediction is out-of-fold (OOF): it is generated by a model that did not train on that observation. This creates genuinely unseen evaluation periods without violating temporal order, which ordinary random k-fold cross-validation would do.",
    )

    heading(doc, "6. Phase 1 methodology and benchmark")
    body(
        doc,
        "Phase 1 built the foundational end-to-end experiment: SEC/FRED ingestion, fiscal-quarter normalization, three central KPIs, company-history and sector-peer features, time-series forecasts, macroeconomic features, regularized logistic regression, and gradient boosting. Five forecast architectures were compared because the best time-series structure differs by KPI and horizon; persistence was treated as a mandatory baseline rather than assuming a complex model would win.",
    )
    body(
        doc,
        "Three expanding development windows were used for model selection. A 2023-and-later block then remained untouched until the champion and threshold were frozen. This is the strongest evidence level in the completed project because it approximates one prospective evaluation.",
    )
    add_table(
        doc,
        ["Phase 1 locked-holdout slice", "N", "PR-AUC", "Recall", "Precision", "Lift", "Brier"],
        [
            ["Overall", "457", "0.397", "0.563", "0.333", "1.97x", "0.159"],
            ["Consumer Discretionary", "228", "0.468", "0.610", "0.439", "2.18x", "0.171"],
            ["Utilities", "229", "0.332", "0.486", "0.225", "2.15x", "0.147"],
        ],
        [2600, 850, 1150, 1150, 1200, 1050, 1360],
    )

    heading(doc, "7. Phase 1 findings")
    body(
        doc,
        "The classifier concentrated events in the high-risk queue, but performance differed materially by sector. Consumer Discretionary was easier to rank even though its median finances were stronger, because the sector contained wider cross-company dispersion and sharper cyclical transitions. Utilities had weaker median coverage, cash flow, and leverage, yet persistent low ratios and regulated capital structures made new transitions harder to distinguish.",
    )
    bullet(
        doc,
        "Consumer median KPIs: 5.25x coverage, 10.12% free-cash-flow margin, and 22.00% debt-to-assets.",
    )
    bullet(
        doc,
        "Utilities median KPIs: 1.81x coverage, -4.00% free-cash-flow margin, and 33.69% debt-to-assets.",
    )
    bullet(
        doc,
        "Forecast features improved the boosted development model in some settings but not consistently across folds or logistic regression; the finding was therefore only partially supported.",
    )
    bullet(
        doc,
        "Several four-quarter forecast intervals were too narrow. A forecast interval is a range intended to contain a stated share of future outcomes; recalibration widens or adjusts that range using prior forecast errors so observed coverage better matches the stated confidence level.",
    )

    heading(doc, "8. Phase 2 motivation and improvements")
    body(
        doc,
        "Phase 2 responds directly to Phase 1's small issuer count, weaker Utility ranking, moderate precision, interval under-coverage, possible distribution shift, and limited model explanations. The objective was not simply to add algorithms. It was to make the analytical claim more defensible, the features easier to explain, and the review policy match the business preference for high recall.",
    )
    add_table(
        doc,
        ["Phase 1 limitation", "Phase 2 response"],
        [
            ["60 active issuers", "Expanded to 117: 75 Consumer Discretionary and 42 Utilities"],
            [
                "Strict optional-field completeness",
                "Core-label reliability retained; optional gaps allowed and handled within folds",
            ],
            [
                "Single pooled architecture",
                "Partially pooled primary, pooled benchmark, constrained boosting challenger",
            ],
            [
                "Moderate recall at frozen threshold",
                "Sector-specific thresholds minimize workload subject to at least 80% recall",
            ],
            [
                "Unstable or overly broad feature set",
                "Fold-local screening, correlation pruning, temporal permutation tests, and stability reporting",
            ],
            [
                "Forecast interval under-coverage",
                "Empirical recalibration by KPI, sector, and horizon",
            ],
            [
                "Limited case interpretation",
                "Plain-language alert explanations, uncertainty indicators, reviews of correct alerts, false alerts, and missed deteriorations, plus prompts for analyst follow-up",
            ],
        ],
        [3350, 6010],
    )
    heading(doc, "9. Feature selection and interpretability")
    body(
        doc,
        "Candidate features were screened inside every outer training fold. The sequence checks missingness and variance, removes highly correlated duplicates, measures whether shuffling a feature repeatedly harms temporal validation performance, and retains variables only when the evidence is sufficiently stable. An ablation is a controlled comparison in which a feature or feature group is removed to determine whether it contributed useful predictive information.",
    )
    body(
        doc,
        "The stable four-quarter features were financially familiar: operating margin; current interest coverage; revenue growth; interest-coverage year-over-year change; interest-coverage volatility; sector-relative interest coverage; and sector-relative free-cash-flow margin. The latest fold also supported interest-coverage trend, cash-flow conversion, and the debt-to-assets trend. These describe profitability, debt-service capacity, direction, volatility, cash generation, leverage, and peer position.",
    )
    body(
        doc,
        "Obscure refinancing features were not forced into the final explanation. Refinancing gap/assets appeared in only one fold and was not stable. Filing delay was removed from the classifier and retained only as a data-quality and case-review field because a later filing can reflect many administrative causes and is not direct evidence of financial stress.",
    )

    heading(doc, "10. Model comparison and interpretability trade-off")
    body(
        doc,
        "Phase 2 compares three model architectures under the same temporal folds to determine how much sector flexibility and nonlinear modeling contribute. The pooled logistic benchmark estimates one common set of financial relationships. The partially pooled logistic model retains those shared relationships while allowing a small number of prespecified Utility-specific differences. This provides sector flexibility without fitting a fully separate Utility model from only 42 issuers, which could produce unstable estimates. Constrained boosting then tests whether nonlinear thresholds and interactions add ranking value beyond the logistic specifications.",
    )
    body(
        doc,
        "Precision and recall depend on where the alert threshold is placed, so the architectures require a common decision point for a fair comparison. The table uses a fixed 20% alert rate: for each model, the highest-risk one-fifth of out-of-fold company-quarter predictions are classified as alerts. Twenty percent was preregistered as a bounded analyst-capacity scenario alongside narrower 5% and 10% queues; it provides a broader reference queue without allowing workload to vary by model. It was not selected to maximize performance and is not the final operating policy. Section 11 instead chooses sector-specific thresholds that satisfy the project's 80% recall requirement.",
    )
    add_table(
        doc,
        ["Phase 2 model", "Role", "PR-AUC", "Recall", "Precision", "Brier"],
        [
            ["Pooled logistic", "Linear benchmark", "0.367", "0.391", "0.415", "0.156"],
            [
                "Partially pooled logistic",
                "Registered interpretable primary",
                "0.379",
                "0.391",
                "0.415",
                "0.157",
            ],
            ["Constrained boosting", "Nonlinear challenger", "0.412", "0.391", "0.415", "0.153"],
        ],
        [2700, 2450, 1050, 1050, 1100, 1010],
    )
    body(
        doc,
        "Partial pooling increased development PR-AUC from 0.367 to 0.379 relative to the pooled benchmark, providing modest evidence that limited sector differences add value. Constrained boosting ranked events best at 0.412 PR-AUC and also produced the lowest Brier score. The partially pooled logistic model nevertheless remains the registered primary because its financial relationships and sector adjustments are easier to explain and audit. The boosting result is retained as transparent evidence of the performance sacrificed for interpretability.",
    )

    doc.add_page_break()
    heading(doc, "11. Recall-first screening policy")
    policy = recall_policy.loc[recall_policy["required_recall"].eq(0.8)].iloc[0]
    body(
        doc,
        f"The operating rule first requires at least 80% recall in each sector and then selects the thresholds that produce the smallest review queue. On the complete Phase 2 development OOF predictions, this produced {policy['overall_recall']:.1%} overall recall, {policy['consumer_recall']:.1%} Consumer recall, {policy['utility_recall']:.1%} Utility recall, {policy['precision']:.1%} precision, and a {policy['alert_rate']:.1%} alert rate.",
    )
    body(
        doc,
        "Consumer and Utility precision or recall are not forced to be equal. Each sector receives its own threshold because base rates, financial structure, and score distributions differ. The minimum recall standard is shared; the achieved values may differ. The high alert rate is a real operating cost, but it follows directly from the stated preference to avoid missed deteriorations in a pre-review tool.",
    )
    body(
        doc,
        "Precision is the share of alerts that satisfy the exact future label. A false positive may still be a useful financially weak-company review, but it consumes capacity and must not be relabeled as a successful prediction. Recall is the share of actual deterioration rows that were alerted.",
    )

    heading(doc, "12. Two-quarter versus four-quarter sensitivity")
    overall = horizon.loc[horizon["sector"].eq("Overall")].set_index("horizon_quarters")
    add_table(
        doc,
        ["Paired OOF result", "4-quarter primary", "2-quarter sensitivity"],
        [
            ["Rows", "397", "397"],
            [
                "Event prevalence",
                f"{overall.loc[4, 'event_prevalence']:.1%}",
                f"{overall.loc[2, 'event_prevalence']:.1%}",
            ],
            ["PR-AUC", f"{overall.loc[4, 'PR_AUC']:.3f}", f"{overall.loc[2, 'PR_AUC']:.3f}"],
            [
                "Recall",
                f"{overall.loc[4, 'recall_at_80pct_policy']:.1%}",
                f"{overall.loc[2, 'recall_at_80pct_policy']:.1%}",
            ],
            [
                "Precision",
                f"{overall.loc[4, 'precision_at_80pct_policy']:.1%}",
                f"{overall.loc[2, 'precision_at_80pct_policy']:.1%}",
            ],
            [
                "Alert rate",
                f"{overall.loc[4, 'alert_rate_at_80pct_policy']:.1%}",
                f"{overall.loc[2, 'alert_rate_at_80pct_policy']:.1%}",
            ],
            [
                "Brier score",
                f"{overall.loc[4, 'Brier_score']:.3f}",
                f"{overall.loc[2, 'Brier_score']:.3f}",
            ],
            [
                "Median warning lead",
                f"{overall.loc[4, 'median_warning_lead_quarters']:.1f} quarters",
                f"{overall.loc[2, 'median_warning_lead_quarters']:.1f} quarters",
            ],
        ],
        [3600, 2880, 2880],
    )
    figure(
        doc,
        FIGURES / "horizon_tradeoff.png",
        "Figure 1. The two-quarter label reduces queue size but loses ranking quality, precision, and warning time.",
    )
    body(
        doc,
        "The two-quarter model retained high recall and lowered the alert rate from 67.3% to 58.9%, but PR-AUC fell by 0.090, precision fell by 6.6 percentage points, and median warning lead fell from two quarters to one. Company-clustered confidence intervals excluded zero for the PR-AUC, precision, and alert-rate changes. The lower two-quarter Brier score partly reflects its lower event prevalence and is not proof of better ranking. Four quarters therefore remains primary; two quarters can be displayed as an optional near-term sensitivity flag.",
    )

    heading(doc, "13. Company-level case studies")
    body(
        doc,
        "How the cases are classified. The outcome is prospective: a company experiences financial deterioration when, within the four fiscal quarters after the decision date, its interest coverage falls below 1.5x and declines by at least 40% from its decision-date level. This realized-outcome rule is separate from the model's sector-specific probability threshold, which determines whether the company enters the analyst review queue before the future outcome is known. A true positive is alerted and later meets the deterioration rule; a false positive is alerted but does not meet both outcome conditions; and a missed deterioration is not alerted but later meets the rule. The ratios below describe only the information available when each prediction was made.",
        "How the cases are classified. ",
    )
    for _, case in cases.iterrows():
        label = {
            "true_positive": "True positive",
            "false_positive": "False positive",
            "false_negative": "Missed deterioration",
        }.get(case["case_type"], str(case["case_type"]).replace("_", " ").title())
        heading(doc, f"{label}: {case['company_name']} ({case['ticker']})", 2)
        body(
            doc,
            f"Decision date {pd.Timestamp(case['decision_at']).date()}. Model probability {case['probability']:.1%}; sector threshold {case['sector_threshold']:.1%}. Interest coverage {case['interest_coverage_ttm']:.2f}x, free-cash-flow margin {case['free_cash_flow_margin_ttm']:.1%}, operating margin {case['operating_margin_ttm']:.1%}, and debt-to-assets {case['total_debt_to_assets']:.1%}.",
        )
        decline = (
            float(case["interest_coverage_ttm"]) - float(case["future_minimum_interest_coverage"])
        ) / abs(float(case["interest_coverage_ttm"]))
        if case["case_type"] == "true_positive":
            body(
                doc,
                f"Why it is a true positive. The probability exceeded the sector threshold, so the model issued an alert. Coverage subsequently fell from {case['interest_coverage_ttm']:.2f}x to {case['future_minimum_interest_coverage']:.2f}x {case['warning_lead_quarters']:.0f} quarters later—a {decline:.1%} decline that crossed below 1.5x and exceeded the 40% decline requirement.",
                "Why it is a true positive. ",
            )
        elif case["case_type"] == "false_positive":
            body(
                doc,
                f"Why it is a false positive. The probability exceeded the sector threshold, so the model issued an alert. However, coverage was already {case['interest_coverage_ttm']:.2f}x at the decision date and its four-quarter minimum was {case['future_minimum_interest_coverage']:.2f}x. Because coverage became less negative rather than declining another 40%, the future outcome did not satisfy both parts of the deterioration rule.",
                "Why it is a false positive. ",
            )
        else:
            body(
                doc,
                f"Why it is a missed deterioration. The {case['probability']:.1%} probability was narrowly below the {case['sector_threshold']:.1%} Utility threshold, so the model did not issue an alert. Coverage subsequently fell from {case['interest_coverage_ttm']:.2f}x to {case['future_minimum_interest_coverage']:.2f}x within {case['warning_lead_quarters']:.0f} quarter—a {decline:.1%} decline that crossed below 1.5x and met the realized-outcome rule.",
                "Why it is a missed deterioration. ",
            )
    heading(doc, "14. Phase 1-to-initial-Phase 2 performance comparison")
    body(
        doc,
        "The initial Phase 2 analysis did not establish a conclusive performance improvement over Phase 1 because the numbers came from different populations and evidence levels. Phase 1's 0.397 PR-AUC is a one-time final holdout result. The initial Phase 2 estimates of 0.379 for the primary model and 0.412 for the challenger are development OOF results after expanding the population and changing the modeling procedure. The challenger exceeded the numerical Phase 1 benchmark, while the interpretable primary was slightly lower, but neither comparison replaced an out-of-time test.",
    )
    add_table(
        doc,
        ["Dimension", "Phase 1", "Initial Phase 2", "Assessment"],
        [
            [
                "Issuer coverage",
                "60",
                "117",
                "Improved breadth within the same two-sector active-company scope",
            ],
            [
                "Final evidence",
                "Locked 2023+ holdout evaluated",
                "Development OOF only at that stage",
                "Later out-of-time evidence is reported in Section 16",
            ],
            [
                "Primary PR-AUC",
                "0.397 holdout",
                "0.379 development OOF",
                "No defensible improvement claim at that stage",
            ],
            [
                "Best challenger PR-AUC",
                "0.397 holdout champion",
                "0.412 development OOF",
                "Promising result that motivated later optimization",
            ],
            [
                "Decision policy",
                "Frozen general threshold; 56.3% recall",
                "Sector thresholds; >=80% recall target",
                "More aligned with analyst-screening objective, with higher workload",
            ],
            [
                "Interpretability/governance",
                "Feature increments and model card",
                "Fold-local selection, partial pooling, calibration, uncertainty, cases, monitoring",
                "Material methodological improvement",
            ],
        ],
        [1900, 2100, 2200, 3160],
    )

    heading(doc, "15. Phase 2 model improvement and ensemble design")
    body(
        doc,
        "The next Phase 2 step held the 117-company population, public-data inputs, four-quarter deterioration outcome, and point-in-time rules constant. This made the experiment a controlled model improvement exercise: performance changes were driven primarily by the modeling and validation process rather than by adding sectors, companies, or a different target.",
    )
    body(
        doc,
        "The comparison covered regularized logistic regression, random forest, histogram gradient boosting, XGBoost, and a nonlinear support-vector challenger. Logistic regression preserved a transparent linear benchmark. Random forest and boosting allowed financial variables to matter through thresholds and interactions—for example, weak cash flow becoming more concerning when leverage is elevated—without requiring those relationships to be specified in advance.",
    )
    body(
        doc,
        "The ensemble experiments included equal and rank averaging, averaging the strongest models, selecting a recent-window winner, performance-weighted averaging, time-adaptive weighting, stacking, and focused blends. The time-adaptive approach used only completed earlier validation windows when assigning weights, so it could not choose a model using the outcome of the period being predicted. Broad averaging and hard window-by-window switching were not consistently better: weaker models diluted strong predictions, and a model that won one period did not necessarily remain strongest in the next.",
    )
    body(
        doc,
        "The selected ensemble assigns 60% weight to pooled XGBoost and 40% to sector-specific XGBoost. The pooled component learns relationships supported across both sectors and benefits the smaller Utility sample. The sector-specific component allows Consumer Discretionary and Utility relationships to differ. The blend therefore balances statistical stability with economically reasonable sector flexibility.",
    )

    heading(doc, "16. Optimized model results and comparison")
    development = optimized["development"]
    sealed = optimized["sealed_late_2024_test"]
    add_table(
        doc,
        ["Model evidence", "ROC-AUC", "PR-AUC", "Recall", "Precision", "Alert rate"],
        [
            [
                "Initial Phase 2 boosting — development OOF",
                "0.716",
                "0.412",
                f"{policy['overall_recall']:.1%}",
                f"{policy['precision']:.1%}",
                f"{policy['alert_rate']:.1%}",
            ],
            [
                "Optimized ensemble — development OOF",
                f"{development['ROC_AUC']:.3f}",
                f"{development['PR_AUC']:.3f}",
                f"{development['recall']:.1%}",
                f"{development['precision']:.1%}",
                f"{development['alert_rate']:.1%}",
            ],
            [
                "Optimized ensemble — sealed late-2024 test",
                f"{sealed['ROC_AUC']:.3f}",
                f"{sealed['PR_AUC']:.3f}",
                f"{sealed['recall']:.1%}",
                f"{sealed['precision']:.1%}",
                f"{sealed['alert_rate']:.1%}",
            ],
        ],
        [3200, 1230, 1230, 1230, 1230, 1240],
    )
    body(
        doc,
        "The development comparison is the clearest measure of improvement because it uses the same expanded population and the same evidence level. ROC-AUC increased from 0.716 to 0.760 and PR-AUC increased from 0.412 to 0.462. At approximately the same 80% recall objective, precision increased from 29.8% to 33.4% and the alert rate declined from 57.6% to 51.3%. The model therefore ranked deteriorations more effectively while sending a smaller share of observations to analysts.",
    )
    body(
        doc,
        "After the model families, feature process, blend weight, and sector thresholds were frozen, the ensemble was evaluated once on decisions from July through December 2024 whose four-quarter outcomes matured during 2025. Across 178 observations from 93 companies and 28 deterioration events, it achieved 0.841 ROC-AUC, 0.494 PR-AUC, and 85.7% recall. ROC-AUC measures how consistently the model ranks deteriorating observations above non-deteriorating observations across all thresholds; PR-AUC focuses more directly on ranking the less common deterioration events.",
    )
    body(
        doc,
        "These results are stronger than the earlier development evidence because they combine higher ranking metrics with a later period that was not used to select the final specification. They do not mean every alert was correct: sealed-test precision was 26.4%, and 51.1% of observations still entered the review queue because the policy intentionally favors finding deteriorations over minimizing false alerts. Company-clustered 95% intervals of 0.742-0.924 for ROC-AUC and 0.317-0.691 for PR-AUC also show that uncertainty remains meaningful.",
    )

    heading(doc, "17. Limitations and governance")
    bullet(
        doc,
        "Potential value of failed-company histories: Adding point-in-time eligible delisted and failed issuers could expose the model to more severe deterioration paths, increase the diversity of observed events, and test whether the financial signals generalize beyond companies that remain active today.",
    )
    bullet(
        doc,
        "Limitations of expanding the population: A valid extension would need to reconstruct historical company universes without selecting firms because they later failed, resolve incomplete filings and identifier changes, and distinguish interest-coverage deterioration from bankruptcy, mergers, acquisitions, and other reasons for delisting. Adding these companies without those controls could introduce look-ahead or selection bias and unintentionally change the outcome being modeled.",
    )
    bullet(
        doc,
        "Sector scope: two sectors do not establish generalization to Industrials, Healthcare, Energy, Financials, or international firms.",
    )
    bullet(
        doc,
        "Utility sample and event counts are smaller; sector results carry greater uncertainty and weaker precision.",
    )
    bullet(
        doc,
        "The interest-coverage label captures one financially meaningful deterioration pattern, not every form of distress.",
    )
    bullet(
        doc,
        "Overlapping future windows make company-quarter rows correlated; company-clustered resampling helps but does not create more independent issuers.",
    )
    bullet(
        doc,
        "Public XBRL facts contain tag heterogeneity, amendments, missing fields, and measurement noise despite normalization controls.",
    )
    bullet(
        doc,
        "Model explanations identify predictive associations, not causes. They cannot substitute for filing, covenant, liquidity, regulatory, and management review.",
    )
    bullet(
        doc,
        "The recall-first policy creates a large review queue. Capacity, review outcomes, and alert fatigue should be monitored prospectively.",
    )
    heading(doc, "18. Bridge analysis and prospective test plan")
    body(
        doc,
        "Phase 2 completed the planned recent-period evaluation by freezing the optimized model before scoring decisions from July through December 2024 against four-quarter outcomes extending through 2025. Information availability, label maturity, and temporal order were preserved. This provides sealed out-of-time evidence for the current experiment, while a later period that matures after the model freeze remains valuable as an additional prospective replication.",
    )
    body(
        doc,
        "The project will also apply the Phase 1 methodology to the expanded 117-company population as part of a controlled Phase 1-Phase 2 bridge analysis. Four specifications will be evaluated using the same four-quarter deterioration label, temporal folds, eligible company-quarter observations, and performance measures:",
    )
    add_table(
        doc,
        ["Test", "Population", "Method", "Purpose"],
        [
            ["A", "Original 60 companies", "Phase 1", "Establish the common-period baseline."],
            [
                "B",
                "Original 60 companies",
                "Phase 2",
                "Compare methodology with A while holding the population constant.",
            ],
            [
                "C",
                "Expanded 117 companies",
                "Phase 1",
                "Measure population expansion under the Phase 1 methodology.",
            ],
            [
                "D",
                "Expanded 117 companies",
                "Phase 2",
                "Compare methodology with C while holding the expanded population constant.",
            ],
        ],
        [1100, 2000, 1600, 4660],
    )
    body(
        doc,
        "This is a useful additional comparison because A versus B and C versus D isolate the effect of changing the methodology, while A versus C and B versus D show how expanding the company population affects each approach. It therefore avoids attributing a performance difference to the model when it may instead result from a different sample.",
    )
    body(
        doc,
        "The bridge analysis will strengthen the connection between the phases, but it will not replace the final prospective test. The current two-sector Phase 2 specification, feature rules, calibration procedure, and sector thresholds should remain frozen. Company quarters after the 31 December 2025 financial cutoff should remain unexamined until their complete four-quarter outcomes are available.",
    )
    bullet(
        doc,
        "Bridge endpoints: four-quarter PR-AUC, Brier score, calibration, recall, precision, and alert workload under both the fixed 20% alert-rate comparison and the recall-first operating policy.",
    )
    bullet(
        doc,
        "Primary prospective endpoint: four-quarter PR-AUC, with Phase 1's 0.397 benchmark treated as historical context rather than a guaranteed target.",
    )
    bullet(
        doc,
        "Operating endpoints: overall and sector recall, precision, alert rate, distinct episode capture, and warning lead.",
    )
    bullet(
        doc,
        "Probability endpoints: Brier score, calibration plots, expected calibration error, and forecast-interval coverage by KPI, sector, and horizon.",
    )
    bullet(
        doc,
        "Uncertainty: company-clustered intervals and results by time period, sector, and issuer concentration.",
    )
    bullet(
        doc,
        "Release rule: report all prespecified results once, including unfavorable results; do not retune thresholds after opening the test.",
    )
    bullet(
        doc,
        "Extension rule: a third sector should be a separately documented research extension after the two-sector baseline is frozen, with a valid taxonomy and adequate issuer/event coverage.",
    )

    heading(doc, "19. Final conclusions")
    body(
        doc,
        "The project achieved its Phase 1 goal: a reproducible, point-in-time financial deterioration prototype with a legitimate locked-holdout benchmark. It also achieved the goals of Phase 2: broader sampling, more defensible missing-data eligibility, interpretable partial pooling, fold-local feature selection, empirical forecast-interval recalibration, recall-first sector thresholds, uncertainty estimates, monitoring, horizon sensitivity, company cases, systematic model optimization, and a sealed out-of-time evaluation.",
    )
    body(
        doc,
        "The optimized ensemble's 0.841 ROC-AUC, 0.494 PR-AUC, and 85.7% recall provide successful later-period test evidence, but one period and two sectors cannot establish universal future performance. The model serves as an additional screening tool that flags companies for further review and helps analysts identify potential interest-coverage problems more efficiently. The most defensible operating design is a four-quarter medium-term primary screen, an optional two-quarter near-term flag, transparent reason codes, and mandatory manual review.",
    )
    body(
        doc,
        "The primary contribution is the integration of financial meaning with experimental discipline: accounting normalization, point-in-time availability, temporal validation, imbalanced learning, calibration, feature stability, interpretable model structure, and explicit governance. That integration is more valuable than presenting one algorithm or one score in isolation.",
    )

    heading(doc, "Appendix A. Reproducibility and evidence map")
    add_table(
        doc,
        ["Artifact", "Role"],
        [
            [
                "README.md",
                "Project status, benchmark, limitations, Phase 2 design, and publication links",
            ],
            [
                "configs/phase2.yml",
                "Frozen population, seed, model, horizon, and threshold policy; tests verify key controls",
            ],
            ["docs/point_in_time_policy.md", "Availability and leakage rules"],
            ["docs/phase2_methodology.md", "Plain-language Phase 2 technical definitions"],
            [
                "reports/generated/phase2_*.csv and reports/source_data/*final_metrics.json",
                "Auditable Phase 2 metrics, feature evidence, monitoring, sensitivity, and sealed-test support",
            ],
        ],
        [3600, 5760],
    )
    trailing_paragraph = doc.paragraphs[-1]
    if not trailing_paragraph.text:
        trailing_paragraph._element.getparent().remove(trailing_paragraph._element)
    return doc


def main() -> None:
    PUBLICATION.mkdir(parents=True, exist_ok=True)
    doc = build_report()
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
