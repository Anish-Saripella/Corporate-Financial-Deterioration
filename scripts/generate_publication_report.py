"""Build the Phase 1 publication-style research report and supporting figures."""

from __future__ import annotations

from pathlib import Path

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
WORKBOOK = ROOT / "outputs/powerbi_stage17/Corporate_Financial_Deterioration_PowerBI_Import.xlsx"
OUT = ROOT / "reports/publication"
FIG = ROOT / "reports/figures/publication"

NAVY = "17365D"
BLUE = "2E75B6"
PALE = "EAF1F8"
GRAY = "667085"
RED = "B42318"
AMBER = "D97706"
GREEN = "067647"


def font(run, size: float = 10.5, bold: bool = False, color: str = "202124", italic: bool = False) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths, strict=True):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            margins = OxmlElement("w:tcMar")
            for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
                node = OxmlElement(f"w:{side}")
                node.set(qn("w:w"), str(value))
                node.set(qn("w:type"), "dxa")
                margins.append(node)
            tc_pr.append(margins)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, value in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade(cell, PALE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        font(p.add_run(value), 9, True, NAVY)
    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            font(p.add_run(value), 8.7)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_body(doc: Document, text: str, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    if bold_lead and text.startswith(bold_lead):
        font(p.add_run(bold_lead), bold=True, color=NAVY)
        text = text[len(bold_lead) :]
    font(p.add_run(text))


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.167
    font(p.add_run(text), 10.3)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    font(p.add_run(text), {1: 16, 2: 13, 3: 12}[level], True, NAVY if level < 3 else BLUE)


def add_figure(doc: Document, path: Path, caption: str, width: float = 6.3) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    font(cap.add_run(caption), 8.5, italic=True, color=GRAY)


def make_figures(watch: pd.DataFrame, model: pd.DataFrame, portfolio: pd.DataFrame, history: pd.DataFrame) -> None:
    FIG.mkdir(parents=True, exist_ok=True)
    plt.rcParams.update({"font.size": 10, "axes.titleweight": "bold", "axes.spines.top": False, "axes.spines.right": False})

    bands = ["Low", "Moderate", "High", "Severe"]
    counts = watch.groupby(["sector", "risk_band"]).size().unstack(fill_value=0).reindex(columns=bands)
    ax = counts.plot(kind="bar", stacked=True, color=["#5B8FF9", "#F6BD16", "#E8684A", "#B42318"], figsize=(8.2, 4.2))
    ax.set(xlabel="", ylabel="Companies")
    ax.set_title("Latest risk-band composition by sector", pad=42)
    ax.tick_params(axis="x", rotation=0)
    ax.legend(ncol=4, frameon=False, loc="upper center", bbox_to_anchor=(0.5, 1.075),
              columnspacing=1.5, handletextpad=0.6)
    fig = ax.get_figure()
    fig.subplots_adjust(top=0.78, bottom=0.16, left=0.10, right=0.98)
    plt.savefig(FIG / "risk_band_by_sector.png", dpi=220); plt.close()

    top = watch.nlargest(10, "probability").sort_values("probability")
    colors = ["#B42318" if x == "Severe" else "#D97706" for x in top["risk_band"]]
    fig, ax = plt.subplots(figsize=(8.2, 4.6))
    ax.barh(top["ticker"], top["probability"] * 100, color=colors)
    ax.set(title="Highest latest deterioration-risk estimates", xlabel="Estimated deterioration risk (%)", ylabel="")
    ax.axvline(watch.loc[watch.alert, "threshold"].median() * 100, color=f"#{NAVY}", ls="--", lw=1.2, label="Median alert threshold")
    ax.legend(frameon=False, loc="lower right")
    plt.tight_layout(); plt.savefig(FIG / "top_risk_companies.png", dpi=220); plt.close()

    metrics = portfolio.set_index("sector")[["median_interest_coverage_x", "median_free_cash_flow_margin", "median_debt_to_assets"]]
    fig, axes = plt.subplots(1, 3, figsize=(8.4, 3.5))
    labels = [("median_interest_coverage_x", "Interest coverage", "x", 1), ("median_free_cash_flow_margin", "FCF margin", "%", 100), ("median_debt_to_assets", "Debt/assets", "%", 100)]
    for ax, (col, title, unit, scale) in zip(axes, labels, strict=True):
        vals = metrics[col] * scale
        ax.bar(["Consumer\nDiscretionary", "Utilities"], vals, color=[f"#{BLUE}", f"#{AMBER}"])
        ax.set_title(title)
        ax.set_ylabel(unit)
        ax.tick_params(axis="x", labelsize=8)
        ax.axhline(0, color="#98A2B3", lw=0.7)
    fig.suptitle("Sector-level financial condition at the latest evaluated snapshot", fontweight="bold")
    plt.tight_layout(); plt.savefig(FIG / "sector_kpi_comparison.png", dpi=220); plt.close()

    hold = model[model["evaluation_sample"].eq("Locked final holdout")].set_index("sector")
    plot = hold.loc[["Consumer Discretionary", "Utilities"], ["PR_AUC", "recall", "precision"]]
    ax = plot.plot(kind="bar", figsize=(8.2, 4.2), color=[f"#{NAVY}", f"#{BLUE}", f"#{AMBER}"])
    ax.set(title="Locked-holdout discrimination and alert-quality metrics", xlabel="", ylabel="Metric value", ylim=(0, 0.7))
    ax.tick_params(axis="x", rotation=0)
    ax.legend(frameon=False, ncol=3, loc="upper center")
    plt.tight_layout(); plt.savefig(FIG / "holdout_sector_performance.png", dpi=220); plt.close()

    champ = model[
        model["is_champion"]
        & model["evaluation_sample"].str.contains("OOF")
        & model["sector"].eq("Overall")
    ].copy()
    champ = champ.sort_values("fold_id")
    fig, ax = plt.subplots(figsize=(8.2, 4.0))
    ax.plot(champ["fold_id"].astype(str), champ["PR_AUC"], marker="o", color=f"#{NAVY}", label="PR-AUC")
    ax.plot(champ["fold_id"].astype(str), champ["Brier_score"], marker="o", color=f"#{AMBER}", label="Brier score")
    ax.set(title="Champion out-of-fold performance across expanding-window folds", xlabel="Temporal fold", ylabel="Metric value")
    ax.legend(frameon=False)
    plt.tight_layout(); plt.savefig(FIG / "temporal_fold_stability.png", dpi=220); plt.close()

    history = history.copy()
    key_parts = history["decision_key"].str.split("|", regex=False)
    history["fiscal_year"] = key_parts.str[1].astype(int)
    history["fiscal_quarter"] = key_parts.str[2]
    trend = (
        history[history["fiscal_year"].between(2012, 2024)]
        .groupby(["fiscal_year", "sector"])[
            ["interest_coverage_ttm", "free_cash_flow_margin_ttm", "total_debt_to_assets"]
        ]
        .median()
        .reset_index()
    )
    fig, axes = plt.subplots(3, 1, figsize=(8.4, 7.0), sharex=True)
    trend_specs = [
        ("interest_coverage_ttm", "Interest coverage", "x", 1),
        ("free_cash_flow_margin_ttm", "Free-cash-flow margin", "%", 100),
        ("total_debt_to_assets", "Debt to assets", "%", 100),
    ]
    colors = {"Consumer Discretionary": f"#{BLUE}", "Utilities": f"#{AMBER}"}
    for ax, (column, title, unit, scale) in zip(axes, trend_specs, strict=True):
        for sector_name, sector_data in trend.groupby("sector"):
            ax.plot(sector_data["fiscal_year"], sector_data[column] * scale, marker="o", ms=3,
                    color=colors[sector_name], label=sector_name)
        ax.set_ylabel(unit); ax.set_title(title, loc="left", fontsize=10)
        ax.axhline(0, color="#98A2B3", lw=0.7)
    axes[0].legend(frameon=False, ncol=2, loc="upper left")
    axes[-1].set_xlabel("Issuer fiscal year")
    fig.suptitle("Sector median KPI trajectories, fiscal years 2012–2024", fontweight="bold")
    plt.tight_layout(); plt.savefig(FIG / "sector_kpi_time_series.png", dpi=220); plt.close()

    seasonal = history.groupby(["fiscal_quarter", "sector"])[
        ["interest_coverage_ttm", "free_cash_flow_margin_ttm", "total_debt_to_assets"]
    ].median().reset_index()
    seasonal["fiscal_quarter"] = pd.Categorical(
        seasonal["fiscal_quarter"], ["FQ1", "FQ2", "FQ3", "FQ4"], ordered=True
    )
    seasonal = seasonal.sort_values("fiscal_quarter")
    fig, axes = plt.subplots(1, 3, figsize=(8.4, 3.3))
    for ax, (column, title, unit, scale) in zip(axes, trend_specs, strict=True):
        for sector_name, sector_data in seasonal.groupby("sector", observed=True):
            ax.plot(sector_data["fiscal_quarter"].astype(str), sector_data[column] * scale,
                    marker="o", color=colors[sector_name])
        ax.set_title(title, fontsize=9); ax.set_ylabel(unit); ax.tick_params(axis="x", labelsize=8)
        ax.axhline(0, color="#98A2B3", lw=0.7)
    fig.suptitle("Full-sample medians by standardized issuer fiscal quarter", fontweight="bold")
    plt.tight_layout(); plt.savefig(FIG / "sector_kpi_fiscal_quarter.png", dpi=220); plt.close()


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5); section.page_height = Inches(11)
    section.top_margin = Inches(1); section.bottom_margin = Inches(1)
    section.left_margin = Inches(1); section.right_margin = Inches(1)
    section.header_distance = Inches(0.492); section.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"; normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.10
    for level, size, before, after in ((1, 16, 16, 8), (2, 13, 12, 6), (3, 12, 8, 4)):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"; style.font.size = Pt(size); style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(NAVY if level < 3 else BLUE)
        style.paragraph_format.space_before = Pt(before); style.paragraph_format.space_after = Pt(after)


def build() -> Path:
    OUT.mkdir(parents=True, exist_ok=True)
    portfolio = pd.read_excel(WORKBOOK, "Portfolio")
    watch = pd.read_excel(WORKBOOK, "Watchlist", dtype={"cik": str})
    model = pd.read_excel(WORKBOOK, "ModelPerformance")
    history = pd.read_excel(WORKBOOK, "CompanyHistory")
    make_figures(watch, model, portfolio, history)
    hold = model[model["evaluation_sample"].eq("Locked final holdout")].set_index("sector")

    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(hp.add_run("CORPORATE FINANCIAL DETERIORATION | PHASE 1 RESEARCH"), 8.5, True, GRAY)
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(fp.add_run("Independent analytical research | Not investment advice or a default-probability model"), 8, color=GRAY)

    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(92); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run("FINANCIAL RISK & DATA SCIENCE RESEARCH"), 10, True, AMBER)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(8)
    font(p.add_run("Corporate Financial Deterioration Early-Warning Platform"), 27, True, NAVY)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(24)
    font(p.add_run("A point-in-time, experimentally governed study of debt-service capacity across Consumer Discretionary and Utilities issuers"), 13, color=BLUE)
    add_table(doc, ["Research design", "Coverage", "Decision use"], [["Expanding-window validation with locked holdout", "60 issuers | 3,150 company-quarters", "Analyst prioritization and surveillance"]], [3120, 3120, 3120])
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(90)
    font(p.add_run("Phase 1 publication report | Data snapshot 10 December 2024 | Prepared 5 August 2026"), 10, color=GRAY)
    doc.add_page_break()

    add_heading(doc, "Executive summary")
    add_body(doc, "Research objective. This study tests whether filing-aware public fundamentals and macroeconomic data can identify weakening debt-service capacity over the subsequent four fiscal quarters. The output is a ranked early-warning estimate for analyst triage; it is explicitly not a bankruptcy, default, valuation, or investment-recommendation model.", "Research objective.")
    add_body(doc, "Experimental design. A frozen universe of 60 currently listed U.S. issuers (30 per sector) was observed at company-fiscal-quarter frequency. SEC fundamentals were normalized across issuer calendars and XBRL concepts, macro variables were joined under historical-availability rules, and all model selection used expanding-window out-of-fold predictions with label-availability embargoes. A 2023+ period was held untouched until the champion specification was frozen.", "Experimental design.")
    add_body(doc, "Principal result. The gradient-boosted champion produced 0.397 PR-AUC, 0.563 recall, 0.333 precision, 1.966 top-decile lift, and 0.159 Brier score on 457 locked-holdout observations. Consumer Discretionary generalized better than Utilities (PR-AUC 0.468 versus 0.332), indicating economically meaningful sector heterogeneity and a need for expanded sector-specific validation.", "Principal result.")
    add_body(doc, "Portfolio finding. At the latest evaluated company snapshot, 22 of 60 firms were alerted, 6 were classified Severe, and mean estimated deterioration risk was 28.13%. Both sectors recorded 11 alerts, but Utilities displayed weaker median interest coverage (2.62x), negative median free-cash-flow margin (-16.39%), and higher debt-to-assets (38.18%) than Consumer Discretionary.", "Portfolio finding.")
    add_body(doc, "Research judgment. The evidence supports use as a transparent screening layer when alerts are paired with company history, forecast uncertainty, sector context, and human review. It does not support automated credit decisions or population-level default inference.", "Research judgment.")
    add_figure(doc, FIG / "risk_band_by_sector.png", "Figure 1. Latest evaluated risk-band distribution. Source: certified Phase 1 analysis workbook.")

    add_heading(doc, "1. Research question and experimental hypotheses")
    add_body(doc, "The primary question is whether changes in interest coverage, cash generation, leverage, peer position, forecasted financial ratios, and economic conditions contain useful information about future debt-service deterioration. Each observation represents one company at one quarterly decision date, using only information available at that time. The study measures how well the model ranks and estimates risk; it does not attempt to prove that any feature causes deterioration.")
    add_table(doc, ["Hypothesis", "Empirical test", "Decision criterion"], [
        ["H1: combining signals improves ranking", "Compare planned feature groups using time-ordered, out-of-sample PR-AUC", "Improvement must persist across time periods and sectors"],
        ["H2: forecasts add forward-looking information", "Add 1Q/4Q KPI forecasts built without future information", "Added value must appear before the final test period"],
        ["H3: performance is regime- and sector-sensitive", "Report fold, sector, calibration, and final-holdout slices", "No pooled metric may conceal weak subgroups"],
    ], [1800, 4200, 3360])
    add_body(doc, "This is an observational prediction study. Feature importance therefore shows which variables helped the model make predictions after considering the other inputs. It does not prove that changing a balance-sheet item would cause the company's risk to rise or fall.")

    add_heading(doc, "2. Data architecture and point-in-time research controls")
    add_body(doc, "Source selection. SEC EDGAR supplies public company financial statements and filing dates, while FRED/ALFRED supplies economic data such as interest rates and business-cycle measures. These public sources make the analysis easier to verify and reproduce at low cost, but they also create data challenges. Companies may use different SEC tags or labels for the same financial concept, filings become available at different times, previously published values may later be revised, and the dataset contains relatively few severely distressed companies. Every download is therefore recorded in a tracking file that identifies the source, download settings, retrieval time, file-integrity check, and software version.")
    add_body(doc, "Panel construction. SEC XBRL tags were mapped to consistent financial concepts and aligned to each company's fiscal quarters. Filing and availability dates were retained so the model could not use a number before it became public. This point-in-time rule is important because joining everything by calendar quarter can accidentally introduce future information. The final history contains 3,150 unique company-quarter decision records.")
    add_body(doc, "Universe controls. Eligibility required coverage of interest coverage, free-cash-flow margin, and debt-to-assets with continuity and denominator checks. Fourteen sampled firms were replaced under prespecified rules. This creates a cleaner experiment but induces survivorship and availability bias; those biases are disclosed rather than hidden in aggregate performance.")
    add_table(doc, ["Analytical layer", "Method choice", "Why the method fits financial data"], [
        ["Accounting normalization", "Concept mapping plus fiscal-period standardization", "Issuer taxonomies and fiscal calendars are nonuniform; economic definitions must be aligned before cross-sectional comparison."],
        ["Point-in-time join", "Use each value only after its public availability date", "Prevents the model from learning from late filings, later revisions, or future outcomes."],
        ["Peer context", "Within-sector percentile ranks", "Financial ratios are structurally sector-dependent; peer ranks reduce misleading level comparisons."],
        ["Missing values", "Preserve valid blanks; estimate replacements within each training window", "Replacing every blank with zero creates false financial values, while using the full dataset would leak future information."],
    ], [1800, 2700, 4860])

    add_heading(doc, "3. Financial constructs, label design, and measurement validity")
    add_body(doc, "Interest coverage is the central debt-service construct because it compares operating earnings with interest burden. Free-cash-flow margin measures internally generated liquidity after investment needs. Debt-to-assets represents balance-sheet leverage and creditor claim intensity. The three measures are complementary: coverage captures service capacity, cash flow captures funding resilience, and leverage captures structural exposure.")
    add_body(doc, "How to interpret the sector comparison. The statement that Utilities are more financially constrained is true within this study's 60-company sample: their median interest coverage is lower, free-cash-flow margin is weaker, debt-to-assets is higher, and short-term liquidity is lower. It should not be generalized into a claim that every utility is less financially healthy than every Consumer Discretionary company. Much of the difference reflects how the industries operate. Utilities must continually fund long-lived generation, transmission, distribution, and environmental infrastructure. These investments create large depreciation charges, sustained capital expenditure, and recurring debt or equity financing needs. Regulated utilities generally recover prudent operating and capital costs through customer rates and earn an allowed return on their rate base, but regulatory approval and cost recovery can occur with a lag. This produces relatively stable demand and revenue while leaving cash flow and interest coverage sensitive to borrowing costs, construction programs, and regulatory timing. Consumer Discretionary businesses generally face less uniform capital requirements and can show stronger coverage and cash generation, but their revenue and margins are more exposed to consumer confidence, employment, inflation, competition, and discretionary spending cycles. In financial-health terms, a Utility's higher leverage or negative free cash flow can be sustainable when supported by predictable regulated cash flows and timely rate recovery; concern increases when leverage rises while coverage falls, financing costs increase, or regulatory recovery is delayed. Conversely, a Consumer Discretionary company can appear stronger on normal-period ratios yet deteriorate faster when demand or margins contract.", "How to interpret the sector comparison.")
    add_body(doc, "A deterioration event occurs when interest coverage falls below 1.5x and declines by at least 40% from its current level within the next four quarters. Requiring both conditions avoids treating a small decline from a strong starting point as a serious warning. Because neighboring quarters share part of the same four-quarter future window, the observations are related over time rather than fully independent.")
    add_figure(doc, FIG / "sector_kpi_comparison.png", "Figure 2. Latest sector medians use separate axes because coverage is a multiple while margin and leverage are percentages.")

    add_heading(doc, "4. Sector exploratory findings: financial structure, trends, and seasonality")
    add_heading(doc, "4.1 Cross-sectional financial profile", level=2)
    add_body(doc, "Consumer Discretionary. Across the full panel, the typical Consumer Discretionary company had 5.25x interest coverage, a 10.12% free-cash-flow margin, and debt equal to 22.00% of assets. These averages describe a sector with greater normal debt-service capacity and internally generated cash than Utilities. The apparent strength is not uniform: 22.5% of company-quarter observations fell below 1.5x coverage, 16.4% had negative free-cash-flow margins, and 27.4% had debt above 40% of assets. The interquartile range in coverage was 11.51x, nearly five times the Utilities spread of 2.32x. Financially, this means the sector contains both highly resilient firms and materially weaker tail cases; the median alone understates issuer-specific risk.", "Consumer Discretionary.")
    add_body(doc, "Utilities. The full-panel Utilities median was 1.81x interest coverage, a -4.00% free-cash-flow margin, and 33.69% debt-to-assets. Coverage was below 1.5x in 43.5% of observations and free cash flow was negative in 55.6%. These levels are consistent with a capital-intensive business model in which large infrastructure investment is frequently financed externally. Negative free cash flow is therefore not automatically evidence of distress. It becomes more informative when paired with falling coverage, rising leverage, weak liquidity, or an unfavorable refinancing environment. Utilities also had stronger clustering around their sector norms, so relatively small changes can matter even when their absolute ratios look less dispersed.", "Utilities.")
    add_table(doc, ["Exploratory indicator", "Consumer Discretionary", "Utilities", "Financial reading"], [
        ["Median interest coverage", "5.25x", "1.81x", "Utilities operate with a thinner recurring debt-service cushion."],
        ["Median FCF margin", "10.12%", "-4.00%", "Utility capital spending often exceeds internally generated cash."],
        ["Median debt/assets", "22.00%", "33.69%", "The regulated asset base generally supports more structural leverage."],
        ["Coverage below 1.5x", "22.5%", "43.5%", "A common absolute threshold carries different base rates by sector."],
        ["Negative FCF margin", "16.4%", "55.6%", "Cash deficits require operating and capital-expenditure context."],
        ["Observed deterioration label", "22.2%", "20.7%", "Similar event rates arise from different financial structures."],
    ], [2100, 1900, 1500, 3860])

    add_heading(doc, "4.2 Time-series and regime findings", level=2)
    add_figure(doc, FIG / "sector_kpi_time_series.png", "Figure 3. Annual medians summarize sector conditions; they do not imply that every issuer followed the median path.")
    add_body(doc, "Consumer Discretionary trajectory. Median coverage declined from 5.42x in fiscal 2012 to 3.57x in 2019, then recovered sharply and peaked at 9.19x in 2022 before moderating to 6.34x in 2024. That improvement did not translate into equally strong cash conversion: median free-cash-flow margin fell from 18.6% in 2019 to 6.0% in 2024, while debt-to-assets rose into 2021–2022 and then eased. The combined reading is more nuanced than the coverage line alone. Post-pandemic earnings recovery strengthened the ability to pay interest, but weakening cash conversion and still-elevated leverage indicate less balance-sheet flexibility than the coverage peak suggests. The deterioration-label rate was 29.3% in 2020–2021 versus 22.7% in 2017–2019, consistent with greater cyclical exposure and uneven issuer recovery.", "Consumer Discretionary trajectory.")
    add_body(doc, "Utilities trajectory. Median coverage rose from 0.85x in 2019 to 3.04x in 2023, then eased to 2.73x in 2024. At the same time, free-cash-flow margin deteriorated from -2.7% to -12.4%, and debt-to-assets increased from 32.1% to 38.9%. The sector therefore shows improving accounting coverage alongside a growing financing requirement. This divergence can occur when operating earnings improve but sustained capital expenditure and debt-funded investment absorb cash. For surveillance, the relevant question is not whether one ratio is good or bad in isolation, but whether earnings, cash generation, leverage, and funding costs remain mutually supportable.", "Utilities trajectory.")
    add_body(doc, "Event interpretation. Consumer Discretionary recorded 92 distinct deterioration episode starts, compared with 101 for Utilities, despite similar row-level label prevalence. Utilities therefore experienced slightly more separate transitions into deterioration, while Consumer Discretionary showed more pronounced regime sensitivity. Because the label requires both sub-1.5x coverage and a 40% decline, higher recent Utility coverage mechanically reduced the recent event rate even as free cash flow and leverage weakened. This is an important limitation of any single outcome definition and supports monitoring the component KPIs alongside the model probability.", "Event interpretation.")

    add_heading(doc, "4.3 Fiscal-quarter patterns and persistence", level=2)
    add_figure(doc, FIG / "sector_kpi_fiscal_quarter.png", "Figure 4. Standardized fiscal-quarter medians. Trailing-twelve-month construction intentionally smooths quarter-specific noise.")
    add_body(doc, "There is little evidence of strong recurring seasonality in the three central KPIs. Consumer Discretionary coverage ranges only from 4.98x to 5.27x across FQ1–FQ4, free-cash-flow margin from 9.1% to 11.3%, and debt-to-assets from 21.7% to 22.4%. Utilities are even steadier: coverage ranges from 1.79x to 1.84x, free-cash-flow margin from -4.2% to -3.8%, and leverage from 33.4% to 34.0%. The modest Consumer Discretionary FQ4 cash-flow decline may reflect working-capital timing or issuer mix, but it is too small to treat as a reliable seasonal trading signal without company-level testing.")
    add_body(doc, "Methodological implication. The weak quarter pattern is expected because the ratios use trailing-twelve-month values, which combine four quarters and reduce seasonal noise. The stronger empirical feature is persistence: recent KPI levels contain substantial information about the next quarter. This explains why a random-walk forecast was difficult to beat for one-quarter cash-flow margin and coverage, while a local-level model was more suitable for gradual leverage changes and four-quarter horizons. Sector-specific baselines, changes from each company's own history, and macro conditions are therefore more defensible than broad seasonal adjustments.", "Methodological implication.")

    add_heading(doc, "5. Forecasting methodology: matching model structure to temporal financial data")
    add_body(doc, "Why time-series methods are necessary. Quarterly fundamentals contain persistence, trend, seasonality-like fiscal effects, structural breaks, irregular volatility, and macro sensitivity. Random train/test splits would mix regimes and overstate generalization. Forecast models therefore use only information available before each origin and are evaluated at fixed 1Q and 4Q horizons.")
    add_table(doc, ["Method", "Financial rationale", "Role in the experiment"], [
        ["Random walk", "Many accounting ratios are persistent and difficult to beat out of sample.", "Mandatory naive benchmark; prevents complexity from receiving credit for persistence."],
        ["Drift", "Allows gradual balance-sheet or margin trajectories.", "Tests whether a simple secular trend adds value over the last observation."],
        ["Local level", "Estimates an underlying financial level while allowing quarterly reporting noise.", "Useful for noisy ratios when the company's true condition changes gradually."],
        ["Local linear trend", "Allows both the financial level and its direction of change to move over time.", "Useful for gradual operating or leverage trends, but can be unstable with short histories."],
        ["Dynamic regression model", "Adds economic variables to a time-series model whose financial level changes over time.", "Tests whether rates and business-cycle conditions improve forecasts."],
    ], [1700, 4000, 3660])
    add_body(doc, "Forecast models are compared using root mean squared error (RMSE), which summarizes the typical size of forecast mistakes while giving larger errors more weight. The study also checks whether forecast ranges contain the eventual result as often as intended. These ranges matter because one point forecast can look more certain than it really is. Several four-quarter ranges were too narrow, so forecast uncertainty remains a limitation.")

    add_heading(doc, "6. Classification methodology and model governance")
    add_body(doc, "Regularized logistic regression provides a linear, directionally interpretable benchmark and controls coefficient instability under correlated ratios. Gradient-boosted trees capture nonlinear thresholds and interactions common in financial deterioration - for example, leverage may be tolerable when coverage and cash generation are strong but hazardous when both weaken. Tree depth, regularization, and feature increments are constrained to reduce variance and preserve auditability.")
    add_body(doc, "How information was added to the model. Before testing began, the research plan specified five groups of information to examine: (1) each company's current and past financial ratios, (2) its position relative to companies in the same sector, (3) forecasts of where its ratios may be heading, built without using future information, (4) economic conditions such as interest rates and business-cycle measures, and (5) a small number of sector-specific combinations.", "How information was added to the model.")
    add_body(doc, "Why test the groups separately? Adding one group at a time shows whether it improves predictions and reduces cherry-picking. Peer comparisons use the correct sector benchmark; forecasts describe direction; economic variables capture shared pressures; and a few sector combinations allow the same signal to matter differently in each industry without testing thousands of unplanned relationships.", "Why test the groups separately?")
    add_body(doc, "Why accuracy alone is misleading. Deterioration occurred in about 21% of observations. Predicting “no deterioration” every time would therefore appear roughly 79% accurate while finding no important cases. PR-AUC measures how well actual deterioration cases are ranked ahead of non-events. Recall is the share of actual cases detected; precision is the share of alerts that became cases. Top-decile lift compares the highest-risk 10% with the portfolio average. Brier score measures probability error, while calibration checks whether predicted risks match observed frequencies. The measures answer different questions and should be reviewed together.", "Why accuracy alone is misleading.")

    add_heading(doc, "7. Temporal validation, leakage prevention, and champion selection")
    add_body(doc, "Three expanding training windows simulate how the model would have been rebuilt and tested as time moved forward. Data preparation, probability adjustment, model fitting, and forecasting are repeated separately inside each window. Outcomes that would not yet have been known are excluded. Out-of-fold means that every development prediction came from a model that was not trained on that observation. Historical rows without a defensible out-of-sample prediction remain unscored.")
    add_body(doc, "The final model was chosen first by development PR-AUC, then by probability reliability, stability across sectors and time, interpretability, and simplicity. The 2023-and-later holdout was kept untouched until the model choice was recorded. Its higher 35.45% alert rate is treated as evidence that the later period differed from development, not as permission to adjust the threshold after seeing the answer.")
    add_figure(doc, FIG / "temporal_fold_stability.png", "Figure 5. Temporal OOF results expose regime sensitivity that a random split would conceal.")

    add_heading(doc, "8. Empirical results")
    add_table(doc, ["Locked-holdout slice", "N", "PR-AUC", "Recall", "Precision", "Lift", "Brier"], [
        [name, f"{int(row.observations):,}", f"{row.PR_AUC:.3f}", f"{row.recall:.3f}", f"{row.precision:.3f}", f"{row.top_decile_lift:.2f}x", f"{row.Brier_score:.3f}"]
        for name, row in hold.loc[["Overall", "Consumer Discretionary", "Utilities"]].iterrows()
    ], [2400, 700, 1100, 1100, 1200, 1100, 1760])
    add_body(doc, "The overall top-decile lift of 1.97x indicates useful concentration of deterioration events in the highest-risk queue. Recall of 56.3% implies that the model identifies more than half of events at the frozen threshold, while precision of 33.3% means approximately one in three alerts corresponds to an observed event under the study definition. This trade-off may be acceptable for low-cost analyst screening but not for automated adverse action.")
    add_body(doc, "Model performance is not the same as sector risk. Consumer Discretionary's higher PR-AUC (0.468 versus 0.332) means the model was better at placing that sector's actual deterioration cases above its non-events. It does not mean that every Consumer Discretionary company is riskier, or that the sector has weaker average finances. The exploratory results show the opposite at the median: Consumer Discretionary has stronger coverage, better free cash flow, and lower leverage. Those median ratios describe the typical company; PR-AUC describes how accurately the model separates the weaker companies from the stronger companies within a sector.", "Model performance is not the same as sector risk.")
    add_body(doc, "Why can stronger median KPIs coexist with deterioration risk? Consumer Discretionary had wider differences between companies and greater sensitivity to economic cycles. Most firms can therefore look financially strong while a smaller group experiences sharp declines. The study's event definition also requires both coverage below 1.5x and a decline of at least 40%, so it measures a transition into serious deterioration rather than a sector's normal financial structure. Across the full history, the event rate was only slightly higher for Consumer Discretionary (22.2%) than Utilities (20.7%). At the latest snapshot, both sectors generated 11 alerts; Consumer Discretionary had six observed events versus four for Utilities and a modestly higher average predicted risk (29.36% versus 26.91%). Utilities had weaker typical ratios, but their ratios were more tightly clustered and often persistently low, making new deterioration cases harder for the model to distinguish. The lower Utilities PR-AUC and precision therefore indicate weaker classification—not proof that Utilities are safer.", "Why can stronger median KPIs coexist with deterioration risk?")
    add_body(doc, "The Utilities Brier score was slightly better despite weaker PR-AUC and precision. In plain language, the average size of its probability errors was smaller, but the model was less successful at ordering the companies that actually deteriorated above those that did not. This is why ranking quality and probability accuracy must be reviewed separately, and why the weaker Utilities result limits how confidently the model should be applied to that sector.")
    add_figure(doc, FIG / "holdout_sector_performance.png", "Figure 6. Locked-holdout sector results. Higher PR-AUC, recall, and precision are better.")

    add_heading(doc, "8.1 Hypothesis findings", level=2)
    add_body(doc, "The three hypotheses introduced at the beginning of the study were evaluated using the time-ordered development tests and the untouched final holdout. The conclusions below distinguish strong evidence from mixed evidence.")
    add_table(doc, ["Hypothesis", "Conclusion", "Evidence and interpretation"], [
        ["H1: Combining signals improves ranking", "Supported", "For gradient-boosted trees, average development PR-AUC increased from 0.332 with current fundamentals alone to 0.416 with macroeconomic variables and limited interactions. The broader feature set exceeded the current-only model in every reported time-period and sector slice. This supports combining company, peer, and economic information for risk ranking."],
        ["H2: Forecasts add forward-looking information", "Partially supported", "For gradient-boosted trees, adding all KPI forecasts increased average development PR-AUC from 0.389 for historical-and-peer information to 0.417. However, the improvement did not appear in every time period, and forecast additions did not improve logistic regression. Forecasts contributed useful information in some settings, but the evidence is not strong enough to claim that they always improve the model."],
        ["H3: Performance changes by sector and time period", "Supported", "Results varied across the three development windows and between sectors. On the final holdout, Consumer Discretionary PR-AUC was 0.468 with 0.439 precision, while Utilities PR-AUC was 0.332 with 0.225 precision. These differences support separate sector and time-period reporting rather than relying only on one overall score."],
    ], [2300, 1500, 5560])
    add_body(doc, "Overall, the study supports the value of combining several types of information and confirms that model performance depends on sector and time period. Forecast features appear promising but require further testing before they should be treated as a consistently reliable improvement.")

    add_heading(doc, "9. Latest portfolio surveillance results")
    add_body(doc, "The latest evaluated watchlist contains 60 unique issuers and reconciles to 22 alerts (36.67%), 6 Severe classifications, and a 28.13% mean risk estimate. Consumer Discretionary and Utilities each contribute 11 alerts, but the underlying economics differ: the Utilities median combines thinner coverage with negative free-cash-flow margin and higher leverage.")
    add_figure(doc, FIG / "top_risk_companies.png", "Figure 7. Highest latest risk estimates; company rankings are analytical screening outputs, not investment recommendations.")
    add_body(doc, "The highest risk names should be investigated through an analyst workflow: review filing recency, reconcile the source facts, inspect KPI and forecast paths, assess one-off accounting items, compare peers, and document whether the signal reflects operating deterioration, financing structure, capital expenditure, or model uncertainty. The risk reason is a diagnostic prompt rather than a causal explanation.")

    add_heading(doc, "10. Limitations and threats to validity")
    for item in [
        "Survivorship and selection bias: the universe contains currently listed firms and excludes delisted/distressed histories, limiting inference about severe credit outcomes.",
        "External validity: two sectors and 60 firms cannot establish robustness across banks, industrials, healthcare, energy, or international accounting regimes.",
        "Outcome construction: the deterioration label is economically motivated but researcher-defined; alternative coverage thresholds or horizons may change event prevalence.",
        "Serial dependence: overlapping four-quarter labels and repeated issuers reduce the effective independent sample size and narrow the set of defensible statistical claims.",
        "Measurement error: XBRL tag diversity, restatements, fiscal-calendar irregularities, denominator sensitivity, and macro vintage proxies can affect features.",
        "Forecast uncertainty: several 4Q interval estimates under-cover, so downstream forecast features should not be treated as exact expected states.",
        "Distribution shift: the locked holdout alert rate and Utilities performance differ from development, indicating temporal and sector instability.",
        "Non-causality: importance and risk reasons describe predictive associations; they do not identify management actions that would cause risk to decline.",
        "Operational validation: no cost-sensitive threshold study, analyst-capacity simulation, prospective shadow deployment, or realized P&L/credit-loss backtest has yet been completed.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "11. Conclusions and practical interpretation")
    add_body(doc, "Phase 1 demonstrates that disciplined data engineering and experimental governance can turn free public financial data into a credible early-warning prototype. The strongest contribution is not a single algorithm; it is the alignment of financial constructs, point-in-time data, time-series forecasting, imbalance-aware classification metrics, temporal validation, and transparent release controls.")
    add_body(doc, "The model shows meaningful ranking value and economically interpretable portfolio signals, especially in Consumer Discretionary. At the same time, moderate precision, sector heterogeneity, forecast-interval under-coverage, and survivorship bias constrain deployment. The appropriate conclusion is controlled analyst augmentation: prioritize review, surface evidence, measure uncertainty, and preserve human judgment.")

    doc.add_page_break()
    add_heading(doc, "Appendix A. Reproducibility and release evidence")
    add_table(doc, ["Control", "Phase 1 implementation"], [
        ["Environment", "Python 3.12 project with locked dependencies and deterministic commands"],
        ["Lineage", "SEC/FRED acquisition manifests, checksums, versioned configurations, and point-in-time policies"],
        ["Model governance", "OOF development predictions, fold-local transformations, frozen champion, one-time holdout"],
        ["Analysis reconciliation", "60 unique companies, 3,150 unique decision records, and 93 model-performance records"],
        ["Release evidence", "Certified analysis workbook, publication figures, source documentation, and automated tests"],
    ], [2600, 6760])
    add_body(doc, "Primary internal references: README.md; docs/point_in_time_policy.md; docs/label_specification.md; docs/model_card.md; docs/modeling_lineage.md; and the stage execution records. External data originate from SEC EDGAR and FRED/ALFRED under their respective public usage policies. Industry interpretation is consistent with U.S. Energy Information Administration descriptions of regulated utility revenue requirements and rate-recovery lags and Federal Energy Regulatory Commission guidance explaining cost-of-service recovery, capital investment, and financing returns.")
    add_body(doc, "Disclaimer. This report is an analytical portfolio project for research and demonstration. It is not investment advice, a credit rating, a bankruptcy/default model, or a substitute for audited financial statements and professional judgment.")

    path = OUT / "Corporate_Financial_Deterioration_Phase1_Research_Report.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    print(build())
